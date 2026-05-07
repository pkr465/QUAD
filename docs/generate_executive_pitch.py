"""
Generate QUAD Executive Pitch Document (.docx)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Colors
QUALCOMM_BLUE = RGBColor(0x32, 0x53, 0xDC)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = "3253DC"
TABLE_ALT_ROW_BG = "F5F7FA"
STAT_BOX_BG = "EEF1FB"


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_margins(cell, top=50, bottom=50, left=100, right=100):
    """Set cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}'
        f' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    # Remove namespace redundancy - just use simple approach
    tcMar2 = parse_xml(f'<w:tcMar {nsdecls("w")}>'
                       f'<w:top w:w="{top}" w:type="dxa"/>'
                       f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
                       f'<w:left w:w="{left}" w:type="dxa"/>'
                       f'<w:right w:w="{right}" w:type="dxa"/>'
                       f'</w:tcMar>')
    tcPr.append(tcMar2)


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def set_table_light_borders(table):
    """Set light gray borders on a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="DDDDDD"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_formatted_paragraph(doc, text, font_name="Calibri", font_size=11,
                            bold=False, italic=False, color=DARK_GRAY,
                            alignment=WD_ALIGN_PARAGRAPH.LEFT,
                            space_before=0, space_after=6):
    """Add a formatted paragraph."""
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return para


def add_section_title(doc, text):
    """Add a section title in Heading 1 style."""
    heading = doc.add_heading(text, level=1)
    for run in heading.runs:
        run.font.name = "Calibri Light"
        run.font.size = Pt(28)
        run.font.color.rgb = QUALCOMM_BLUE
        run.font.bold = True
    heading.paragraph_format.space_after = Pt(24)
    return heading


def add_bullet_point(doc, title, description, space_after=12):
    """Add a bullet point with bold title and description."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.left_indent = Cm(1)

    run_title = para.add_run(title + ": ")
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(11)
    run_title.font.bold = True
    run_title.font.color.rgb = DARK_GRAY

    run_desc = para.add_run(description)
    run_desc.font.name = "Calibri"
    run_desc.font.size = Pt(11)
    run_desc.font.color.rgb = MEDIUM_GRAY
    return para


def create_table_with_style(doc, rows, cols, header_row=True):
    """Create a styled table."""
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_light_borders(table)
    return table


def style_table_cell(cell, text, bold=False, color=DARK_GRAY, font_size=10,
                     alignment=WD_ALIGN_PARAGRAPH.LEFT, bg_color=None):
    """Style a table cell."""
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = alignment
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = color
    if bg_color:
        set_cell_shading(cell, bg_color)


def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()


# ============================================================
# CREATE DOCUMENT
# ============================================================

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = DARK_GRAY

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ============================================================
# PAGE 1: TITLE
# ============================================================

# Add vertical spacing
for _ in range(4):
    doc.add_paragraph().paragraph_format.space_after = Pt(24)

# QUAD title
add_formatted_paragraph(doc, "QUAD", font_name="Calibri Light", font_size=48,
                        bold=True, color=QUALCOMM_BLUE,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

# Subtitle
add_formatted_paragraph(doc, "Qualcomm Unified Agent for Developers",
                        font_name="Calibri Light", font_size=20,
                        color=DARK_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        space_after=18)

# Tagline
add_formatted_paragraph(doc, "The CUDA of Qualcomm",
                        font_name="Calibri Light", font_size=16,
                        italic=True, color=LIGHT_GRAY,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

# Separator line
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para.paragraph_format.space_after = Pt(24)
run = para.add_run("_" * 60)
run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
run.font.size = Pt(8)

# Date and classification
add_formatted_paragraph(doc, "April 2026  |  Qualcomm Internal",
                        font_size=11, color=MEDIUM_GRAY,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

# Mission statement
add_formatted_paragraph(doc,
    "Making every AI developer productive on Qualcomm silicon in < 5 minutes",
    font_name="Calibri Light", font_size=14, italic=True,
    color=QUALCOMM_BLUE, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_page_break(doc)

# ============================================================
# PAGE 2: THE PROBLEM
# ============================================================

add_section_title(doc, "1. Developer Friction is Killing Adoption")

add_formatted_paragraph(doc, "", space_after=12)

add_bullet_point(doc, "SDK Sprawl",
    "5 separate SDKs (QNN, SNPE, Hexagon, Adreno, AIMET) \u2014 each with its own download, setup process, and API surface. Developers must learn all five.")

add_bullet_point(doc, "Manual Profiling",
    "Developers must launch 3+ separate profiling tools to understand performance. No unified view of CPU, GPU, and NPU utilization.")

add_bullet_point(doc, "No Unified Model",
    "Different workflows for every chipset and platform combination. Windows, Android, and IoT each require completely different approaches.")

add_bullet_point(doc, "High Barrier to Entry",
    "Time to First Inference (TTFI) currently takes DAYS due to environment setup alone. Developers abandon Qualcomm before running their first model.")

# Stat box
add_formatted_paragraph(doc, "", space_after=24)
table = doc.add_table(rows=1, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table.cell(0, 0)
set_cell_shading(cell, STAT_BOX_BG)
cell.text = ""
para = cell.paragraphs[0]
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para.paragraph_format.space_before = Pt(16)
para.paragraph_format.space_after = Pt(16)
run = para.add_run("Current TTFI: Several Days")
run.font.name = "Calibri"
run.font.size = Pt(14)
run.font.color.rgb = MEDIUM_GRAY
run = para.add_run("    \u2192    ")
run.font.size = Pt(14)
run.font.color.rgb = DARK_GRAY
run = para.add_run("Target: < 5 Minutes")
run.font.name = "Calibri"
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = QUALCOMM_BLUE

# Set table width
for cell in table.columns[0].cells:
    cell.width = Cm(16)

remove_table_borders(table)

add_page_break(doc)

# ============================================================
# PAGE 3: THE SOLUTION
# ============================================================

add_section_title(doc, "2. QUAD: One Platform. All Qualcomm Silicon.")

add_formatted_paragraph(doc, "Unified Architecture Stack", font_size=13,
                        bold=True, color=DARK_GRAY, space_before=12, space_after=18)

# Architecture table
layers = [
    ("Layer 7", "CLI + IDE Plugins + AI Agent (MCP)", "Developer Interface"),
    ("Layer 6", "Inference Server (QUAD Serve)", "Deployment"),
    ("Layer 5", "Optimizer (Graph Fusion + Quantization)", "Optimization"),
    ("Layer 4", "Libraries (QualcommDNN + QualcommBLAS)", "Math & AI Ops"),
    ("Layer 3", "Runtime (Device + Tensor + Model + Stream)", "Execution"),
    ("Layer 2", "Compiler (QUAD IR + Portable Binary)", "Compilation"),
    ("Layer 1", "Hardware (CPU + GPU + NPU)", "Silicon"),
]

table = doc.add_table(rows=len(layers) + 1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_light_borders(table)

# Header
headers = ["Layer", "Component", "Purpose"]
for i, h in enumerate(headers):
    style_table_cell(table.rows[0].cells[i], h, bold=True, color=WHITE,
                     font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                     bg_color=TABLE_HEADER_BG)

# Data rows
for row_idx, (layer, component, purpose) in enumerate(layers):
    bg = TABLE_ALT_ROW_BG if row_idx % 2 == 0 else None
    style_table_cell(table.rows[row_idx + 1].cells[0], layer, bold=True,
                     color=QUALCOMM_BLUE, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                     bg_color=bg)
    style_table_cell(table.rows[row_idx + 1].cells[1], component,
                     bg_color=bg)
    style_table_cell(table.rows[row_idx + 1].cells[2], purpose,
                     color=MEDIUM_GRAY, bg_color=bg)

# Set column widths
for cell in table.columns[0].cells:
    cell.width = Cm(2.5)
for cell in table.columns[1].cells:
    cell.width = Cm(10)
for cell in table.columns[2].cells:
    cell.width = Cm(4)

add_formatted_paragraph(doc, "", space_after=18)
add_formatted_paragraph(doc,
    "One API across all layers. Mock-first development. Real hardware when ready.",
    font_size=12, italic=True, color=QUALCOMM_BLUE,
    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_page_break(doc)

# ============================================================
# PAGE 4: WHY NOW
# ============================================================

add_section_title(doc, "3. The Market Window is Open")

add_formatted_paragraph(doc, "", space_after=12)

add_bullet_point(doc, "NPU Performance Breakthrough",
    "NPUs now exceed 45 TOPS \u2014 competitive with data center GPUs for inference workloads. Edge hardware is finally ready for production AI.")

add_bullet_point(doc, "On-Device AI Explosion",
    "Privacy requirements, latency demands, and offline scenarios are driving AI to the edge. Qualcomm owns mobile and edge silicon.")

add_bullet_point(doc, "NVIDIA's Blind Spot",
    "NVIDIA has ZERO mobile or edge presence \u2014 yet developers default to CUDA because there is no alternative platform. QUAD fills this gap.")

# Quote box
add_formatted_paragraph(doc, "", space_after=24)
table = doc.add_table(rows=1, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table.cell(0, 0)
set_cell_shading(cell, STAT_BOX_BG)
cell.text = ""
para = cell.paragraphs[0]
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para.paragraph_format.space_before = Pt(18)
para.paragraph_format.space_after = Pt(18)
run = para.add_run("Qualcomm ships ")
run.font.name = "Calibri"
run.font.size = Pt(13)
run.font.color.rgb = DARK_GRAY
run = para.add_run("3+ BILLION")
run.font.name = "Calibri"
run.font.size = Pt(13)
run.font.bold = True
run.font.color.rgb = QUALCOMM_BLUE
run = para.add_run(" SoCs/year. CUDA reaches millions. ")
run.font.name = "Calibri"
run.font.size = Pt(13)
run.font.color.rgb = DARK_GRAY
run = para.add_run("QUAD reaches billions.")
run.font.name = "Calibri"
run.font.size = Pt(13)
run.font.bold = True
run.font.color.rgb = QUALCOMM_BLUE

for cell in table.columns[0].cells:
    cell.width = Cm(16)
remove_table_borders(table)

add_page_break(doc)

# ============================================================
# PAGE 5: WHAT WE BUILT
# ============================================================

add_section_title(doc, "4. Platform Status: Feature-Complete (Mock Mode)")

add_formatted_paragraph(doc, "", space_after=12)

# Stats table
stats_table = doc.add_table(rows=1, cols=4)
stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_borders(stats_table)

stats = [("933", "Tests Passing"), ("15", "Modules"), ("30+", "Templates"), ("15,000+", "Lines of Code")]
for i, (num, label) in enumerate(stats):
    cell = stats_table.rows[0].cells[i]
    set_cell_shading(cell, STAT_BOX_BG)
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(num)
    run.font.name = "Calibri Light"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = QUALCOMM_BLUE

    para2 = cell.add_paragraph()
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para2.paragraph_format.space_after = Pt(12)
    run2 = para2.add_run(label)
    run2.font.name = "Calibri"
    run2.font.size = Pt(9)
    run2.font.color.rgb = MEDIUM_GRAY

add_formatted_paragraph(doc, "", space_after=18)
add_formatted_paragraph(doc, "Module Architecture", font_size=13,
                        bold=True, color=DARK_GRAY, space_after=12)

# Module grid (3x5)
modules = [
    ["Runtime", "Compiler", "Libraries", "Optimizer", "Profiler"],
    ["Kernels", "Serve", "CLI", "UDO", "PSNPE"],
    ["Adapters", "Platforms", "Codegen", "Models", "Utils"],
]

mod_table = doc.add_table(rows=3, cols=5)
mod_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_light_borders(mod_table)

for r_idx, row_data in enumerate(modules):
    for c_idx, mod_name in enumerate(row_data):
        cell = mod_table.rows[r_idx].cells[c_idx]
        bg = TABLE_ALT_ROW_BG if r_idx % 2 == 0 else None
        style_table_cell(cell, mod_name, bold=True, color=QUALCOMM_BLUE,
                         font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         bg_color=bg)

add_formatted_paragraph(doc, "", space_after=18)
add_formatted_paragraph(doc,
    "All modules working in mock mode. Real SDK adapter ready \u2014 blocked only on SDK CLI documentation access.",
    font_size=11, italic=True, color=MEDIUM_GRAY,
    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_page_break(doc)

# ============================================================
# PAGE 6: COMPETITIVE ADVANTAGE VS CUDA
# ============================================================

add_section_title(doc, "5. Where QUAD Wins")

add_formatted_paragraph(doc, "", space_after=12)

# Comparison table
comp_data = [
    ["Dimension", "CUDA", "QUAD"],
    ["Market", "Data center / HPC", "Edge / Mobile / AI PC / IoT"],
    ["Compute", "GPU only", "CPU + GPU + NPU (heterogeneous)"],
    ["Power", "300\u2013700W", "1\u201315W (power as a feature)"],
    ["Devices", "Millions of GPUs", "Billions of SoCs"],
    ["Differentiator", "Raw throughput", "Performance-per-watt + AI agent"],
    ["Developer UX", "Manual SDK", "Natural language (MCP)"],
]

comp_table = doc.add_table(rows=len(comp_data), cols=3)
comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_light_borders(comp_table)

# Header row
for i, h in enumerate(comp_data[0]):
    style_table_cell(comp_table.rows[0].cells[i], h, bold=True, color=WHITE,
                     font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                     bg_color=TABLE_HEADER_BG)

# Data rows
for r_idx in range(1, len(comp_data)):
    bg = TABLE_ALT_ROW_BG if r_idx % 2 == 1 else None
    style_table_cell(comp_table.rows[r_idx].cells[0], comp_data[r_idx][0],
                     bold=True, color=DARK_GRAY, bg_color=bg)
    style_table_cell(comp_table.rows[r_idx].cells[1], comp_data[r_idx][1],
                     color=MEDIUM_GRAY, bg_color=bg)
    style_table_cell(comp_table.rows[r_idx].cells[2], comp_data[r_idx][2],
                     bold=True, color=QUALCOMM_BLUE, bg_color=bg)

# Set column widths
for cell in comp_table.columns[0].cells:
    cell.width = Cm(3.5)
for cell in comp_table.columns[1].cells:
    cell.width = Cm(6)
for cell in comp_table.columns[2].cells:
    cell.width = Cm(7)

add_page_break(doc)

# ============================================================
# PAGE 7: QUAD DIFFERENTIATORS
# ============================================================

add_section_title(doc, "6. What CUDA Can Never Do")

add_formatted_paragraph(doc, "", space_after=12)

# 4 differentiator boxes using a 2x2 table
diff_data = [
    ("Power-Aware Computing", "Every API accepts power budget as a constraint. Developers specify watt limits, and QUAD optimizes execution across CPU, GPU, and NPU to stay within thermal and power envelopes."),
    ("AI Agent (MCP)", "Natural language to production code in seconds. Developers describe what they want; the agent generates optimized, hardware-aware code automatically."),
    ("Heterogeneous Orchestration", "CPU + GPU + NPU allocation in a single API call. QUAD schedules work across all compute units based on workload characteristics and power constraints."),
    ("Edge-Native Design", "Built from the ground up for 1\u201315W devices with thermal-aware scheduling, battery optimization, and real-time inference at the edge."),
]

for i, (title, desc) in enumerate(diff_data):
    # Create a single-cell table as a "box"
    box_table = doc.add_table(rows=1, cols=1)
    box_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = box_table.cell(0, 0)
    set_cell_shading(cell, TABLE_ALT_ROW_BG)
    cell.text = ""

    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(f"{i+1}. {title}")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = QUALCOMM_BLUE

    para2 = cell.add_paragraph()
    para2.paragraph_format.space_after = Pt(10)
    run2 = para2.add_run(desc)
    run2.font.name = "Calibri"
    run2.font.size = Pt(10)
    run2.font.color.rgb = MEDIUM_GRAY

    for c in box_table.columns[0].cells:
        c.width = Cm(16)

    # Remove borders, use shading only
    remove_table_borders(box_table)

    # Space between boxes
    add_formatted_paragraph(doc, "", space_after=8)

add_page_break(doc)

# ============================================================
# PAGE 8: SUCCESS METRICS & TIMELINE
# ============================================================

add_section_title(doc, "7. Measurable Impact")

add_formatted_paragraph(doc, "", space_after=12)

# Metrics table
metrics_data = [
    ["Metric", "Target", "Timeline"],
    ["TTFI (new developer)", "< 5 minutes", "Q3 2026"],
    ["Developer NPS", "\u2265 60", "Q4 2026"],
    ["pip install \u2192 inference", "Works out of box", "Q2 2026"],
    ["Community SO questions", "> 1,000", "Q4 2026"],
    ["Pre-optimized model zoo", "100+ models", "Q3 2026"],
    ["Enterprise deployments", "\u2265 10", "Q1 2027"],
]

met_table = doc.add_table(rows=len(metrics_data), cols=3)
met_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_light_borders(met_table)

for i, h in enumerate(metrics_data[0]):
    style_table_cell(met_table.rows[0].cells[i], h, bold=True, color=WHITE,
                     font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                     bg_color=TABLE_HEADER_BG)

for r_idx in range(1, len(metrics_data)):
    bg = TABLE_ALT_ROW_BG if r_idx % 2 == 1 else None
    style_table_cell(met_table.rows[r_idx].cells[0], metrics_data[r_idx][0],
                     bold=True, color=DARK_GRAY, bg_color=bg)
    style_table_cell(met_table.rows[r_idx].cells[1], metrics_data[r_idx][1],
                     color=QUALCOMM_BLUE, bold=True,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, bg_color=bg)
    style_table_cell(met_table.rows[r_idx].cells[2], metrics_data[r_idx][2],
                     color=MEDIUM_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                     bg_color=bg)

for cell in met_table.columns[0].cells:
    cell.width = Cm(6)
for cell in met_table.columns[1].cells:
    cell.width = Cm(5)
for cell in met_table.columns[2].cells:
    cell.width = Cm(4)

# Timeline section
add_formatted_paragraph(doc, "", space_after=24)
add_formatted_paragraph(doc, "Delivery Timeline", font_size=13,
                        bold=True, color=DARK_GRAY, space_after=12)

# Timeline as a flow
timeline_phases = [
    ("Phase B", "Runtime"),
    ("Phase C", "Libraries"),
    ("Phase D", "Profiler"),
    ("Phase E", "Kernels"),
    ("Phase F", "Serve"),
    ("Phase G", "Ecosystem"),
]

tl_table = doc.add_table(rows=2, cols=6)
tl_table.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_borders(tl_table)

for i, (phase, name) in enumerate(timeline_phases):
    cell_top = tl_table.rows[0].cells[i]
    set_cell_shading(cell_top, TABLE_HEADER_BG)
    style_table_cell(cell_top, phase, bold=True, color=WHITE,
                     font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                     bg_color=TABLE_HEADER_BG)

    cell_bot = tl_table.rows[1].cells[i]
    set_cell_shading(cell_bot, TABLE_ALT_ROW_BG)
    style_table_cell(cell_bot, name, color=DARK_GRAY, font_size=9,
                     alignment=WD_ALIGN_PARAGRAPH.CENTER,
                     bg_color=TABLE_ALT_ROW_BG)

add_formatted_paragraph(doc, "", space_after=12)
add_formatted_paragraph(doc,
    "All phases COMPLETE in mock mode. Real hardware integration: Q3\u2013Q4 2026.",
    font_size=11, italic=True, color=MEDIUM_GRAY,
    alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_page_break(doc)

# ============================================================
# PAGE 9: THE ASK
# ============================================================

add_section_title(doc, "8. What We Need")

add_formatted_paragraph(doc, "", space_after=18)

# 3 asks
asks = [
    ("1. QNN SDK Documentation Access",
     "Unblocks the real hardware adapter in 2\u20133 days. The mock framework is complete; we need only the official API surface to wire in real inference."),
    ("2. Hardware: 3 Target Devices",
     "Snapdragon X Elite (AI PC), Arduino UNO Q (IoT), and Snapdragon 8 Elite (Mobile). These cover all three target segments."),
    ("3. Engineering Team: 12\u201315 Engineers over 12 Months",
     "To achieve CUDA platform parity across compiler, runtime, libraries, and ecosystem tooling."),
]

for title, desc in asks:
    box_table = doc.add_table(rows=1, cols=1)
    box_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = box_table.cell(0, 0)
    set_cell_shading(cell, TABLE_ALT_ROW_BG)
    cell.text = ""

    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = QUALCOMM_BLUE

    para2 = cell.add_paragraph()
    para2.paragraph_format.space_after = Pt(12)
    run2 = para2.add_run(desc)
    run2.font.name = "Calibri"
    run2.font.size = Pt(11)
    run2.font.color.rgb = MEDIUM_GRAY

    for c in box_table.columns[0].cells:
        c.width = Cm(16)
    remove_table_borders(box_table)
    add_formatted_paragraph(doc, "", space_after=8)

# Phasing note
add_formatted_paragraph(doc, "", space_after=12)
add_formatted_paragraph(doc,
    "Phase 1 (current team) \u2192 Phase 2 (expand) \u2192 Phase 3 (community & ecosystem)",
    font_size=11, bold=True, color=DARK_GRAY,
    alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_page_break(doc)

# ============================================================
# PAGE 10: CALL TO ACTION
# ============================================================

add_section_title(doc, "9. The Opportunity")

# Spacer
for _ in range(2):
    add_formatted_paragraph(doc, "", space_after=18)

add_formatted_paragraph(doc,
    "3 billion Qualcomm SoCs shipped per year.",
    font_name="Calibri Light", font_size=16, color=DARK_GRAY,
    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

add_formatted_paragraph(doc,
    "Zero unified developer platform.",
    font_name="Calibri Light", font_size=16, color=DARK_GRAY,
    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

add_formatted_paragraph(doc,
    "QUAD makes Qualcomm the default choice for on-device AI development.",
    font_name="Calibri Light", font_size=14, italic=True, color=QUALCOMM_BLUE,
    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

# Command box
cmd_table = doc.add_table(rows=1, cols=1)
cmd_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = cmd_table.cell(0, 0)
set_cell_shading(cell, "1A1A2E")
cell.text = ""
para = cell.paragraphs[0]
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para.paragraph_format.space_before = Pt(16)
para.paragraph_format.space_after = Pt(16)
run = para.add_run("pip install qualcomm-ai-toolkit && quad quickstart")
run.font.name = "Calibri"
run.font.size = Pt(13)
run.font.bold = True
run.font.color.rgb = RGBColor(0x6C, 0xE0, 0x6C)

for c in cmd_table.columns[0].cells:
    c.width = Cm(14)
remove_table_borders(cmd_table)

add_formatted_paragraph(doc, "", space_after=24)

add_formatted_paragraph(doc,
    "Demo available:  ./install.sh && ./launch.sh",
    font_size=12, color=MEDIUM_GRAY,
    alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

# Final separator
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
para.paragraph_format.space_after = Pt(18)
run = para.add_run("_" * 60)
run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
run.font.size = Pt(8)

add_formatted_paragraph(doc,
    "QUAD  |  Qualcomm Internal  |  April 2026",
    font_size=10, color=LIGHT_GRAY,
    alignment=WD_ALIGN_PARAGRAPH.CENTER)

# ============================================================
# SAVE
# ============================================================

output_path = "/Users/pavanr/work/05/QUAD/docs/QUAD_Executive_Pitch.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
print(f"File size: {os.path.getsize(output_path) / 1024:.1f} KB")
