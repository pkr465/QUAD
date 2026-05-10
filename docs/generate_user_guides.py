"""Generate the Server-side and Client-side QUAD user guides as .docx.

Produces two files:
    docs/QUAD_Server_Guide.docx   (operator / DevOps audience)
    ../QUAD-Client/docs/QUAD_Client_Guide.docx   (developer audience)

Run from anywhere:
    python docs/generate_user_guides.py

Requires python-docx (`pip install python-docx`). All content is inline
in this script — no external Markdown dependency — so the docx files
can be regenerated whenever the source repo's reality changes.
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ─── Styling helpers ───────────────────────────────────────────────────────


def _setup_styles(doc: Document) -> None:
    """Tweak default Word styles so the guide renders consistently."""
    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(11)

    for h, sz in (("Heading 1", 22), ("Heading 2", 16), ("Heading 3", 13)):
        st = doc.styles[h]
        st.font.name = "Calibri"
        st.font.size = Pt(sz)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5C)

    # Custom code-block style (monospace + light-grey background)
    if "QuadCode" not in [s.name for s in doc.styles]:
        cstyle = doc.styles.add_style("QuadCode", WD_STYLE_TYPE.PARAGRAPH)
        cstyle.font.name = "Consolas"
        cstyle.font.size = Pt(9.5)
        cstyle.font.color.rgb = RGBColor(0x10, 0x10, 0x10)
        cstyle.paragraph_format.space_before = Pt(2)
        cstyle.paragraph_format.space_after = Pt(2)
        cstyle.paragraph_format.left_indent = Inches(0.15)


def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_paragraph_shading(par, hex_color: str) -> None:
    p_pr = par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    p_pr.append(shd)


def code(doc: Document, text: str) -> None:
    """Render a code block — monospace, grey background."""
    for line in text.splitlines() or [""]:
        p = doc.add_paragraph(style="QuadCode")
        p.add_run(line)
        _add_paragraph_shading(p, "F2F2F2")


def callout(doc: Document, text: str, kind: str = "note") -> None:
    """Inline callout box: note / warning / tip."""
    palette = {
        "note":    ("E8F0FE", "1A73E8", "Note"),
        "warning": ("FCE8E6", "C5221F", "Warning"),
        "tip":     ("E6F4EA", "188038", "Tip"),
    }
    bg, fg, label = palette.get(kind, palette["note"])
    p = doc.add_paragraph()
    run = p.add_run(f"  {label}: ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(fg)
    p.add_run(text)
    _add_paragraph_shading(p, bg)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Render a borderless data table with a coloured header row."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.autofit = True
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
        _set_cell_shading(cell, "1F3A5C")
        for r in cell.paragraphs[0].runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = val


def bullets(doc: Document, items: Iterable[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def numbered(doc: Document, items: Iterable[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Number")


def title_page(doc: Document, title: str, subtitle: str, audience: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(36)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5C)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.italic = True
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(40)
    r3 = p3.add_run(f"Audience: {audience}")
    r3.font.size = Pt(11)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("QUAD v0.4.0")
    r4.font.size = Pt(10)
    r4.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_page_break()


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def p(doc, text):
    doc.add_paragraph(text)


# ═══════════════════════════════════════════════════════════════════════════
# SERVER GUIDE
# ═══════════════════════════════════════════════════════════════════════════


def build_server_guide(out_path: Path) -> None:
    doc = Document()
    _setup_styles(doc)

    title_page(
        doc,
        "QUAD Server Guide",
        "Operator playbook for the Qualcomm Unified Agent for Developers MCP server",
        "Server operators, DevOps engineers, AI infrastructure teams",
    )

    # ─── 1. What is QUAD ──────────────────────────────────────────────────
    h1(doc, "1. What is QUAD?")
    p(doc, "QUAD (Qualcomm Unified Agent for Developers) is an MCP server that "
           "exposes Qualcomm's QAIRT / SNPE / QNN SDKs to AI developers through "
           "natural language. Instead of memorising 18 CLI tools across 6 OS "
           "platforms, developers ask Claude Code (or any MCP-compatible IDE) "
           "to convert a model, profile it, allocate ops across CPU/GPU/NPU, "
           "or generate runtime code — and QUAD invokes the right SDK tool for them.")
    p(doc, "This guide is for the people running the QUAD server: "
           "installing it, keeping the SDK current, monitoring health, "
           "scaling for production, and shipping releases.")

    h2(doc, "1.1  Goals")
    bullets(doc, [
        "Time-to-First-Inference < 10 minutes from a clean machine",
        "Mock-mode CI without any SDK installed",
        "Real-mode end-to-end on Snapdragon X-series Copilot+ PCs",
        "Single command upgrades (./install.sh re-runs are idempotent)",
    ])

    h2(doc, "1.2  When to use QUAD vs. native SDK tools")
    table(doc, ["Situation", "Use QUAD", "Use the SDK tools directly"], [
        ["Iterating on model graphs / quantization choices", "Yes — fastest loop", "No"],
        ["Production inference at scale (>10 QPS)", "Yes (HTTP server)", "Either"],
        ["Custom op packages, deep DSP kernel work", "No — bypass adapter", "Yes"],
        ["Shipping a wheel that ships to other developers", "Yes (quad-agent on PyPI)", "No"],
    ])

    # ─── 2. Architecture ──────────────────────────────────────────────────
    h1(doc, "2. Architecture")
    p(doc, "QUAD has two cleanly separated pieces:")
    bullets(doc, [
        "Server (~1.2 GB with QAIRT) — the heavy bits: SDK adapter, compiler, codegen templates, AIMET / AI Hub integrations. This is what this guide covers.",
        "Client (~6 MB) — a lightweight provisioner that wires Claude Code (or any MCP-compatible IDE) to the server. See QUAD_Client_Guide.docx.",
    ])

    h2(doc, "2.1  Seven-layer stack")
    code(doc, """
+-------------------------------------------------------------------+
|  Layer 7: DevX / Ecosystem                                         |
|  CLI (mode/sdk/doctor/quickstart)  VS Code Extension  Plugins     |
+-------------------------------------------------------------------+
|  Layer 6: Serve & Deploy                                           |
|  ModelServer  ModelRegistry  deploy_model()  FastAPI HTTP server  |
+-------------------------------------------------------------------+
|  Layer 5: MCP Agent (5 tools via FastMCP)                          |
|  hardware_detect  convert_model  profile_workload                 |
|  orchestrate_workload  generate_code                              |
|  + sdk_manager.startup_resolve_and_log on every start             |
+-------------------------------------------------------------------+
|  Layer 4: SDK Abstraction (Adapter Pattern)                        |
|  MockAdapter  <->  QAIRTAdapter   [config: mock | real]           |
|  + AdapterFactory.strict mode  (QUAD_STRICT_REAL=1)               |
|  AIMET, AI Hub, Hexagon, Adreno wrappers                          |
+-------------------------------------------------------------------+
|  Layer 3: Libraries & Optimizer                                    |
|  QualcommDNN (Conv, MHA, Flash)  QualcommBLAS (GEMM)              |
|  Fusion / DCE / ConstantFold / MemoryPlanning                     |
+-------------------------------------------------------------------+
|  Layer 2: Runtime & Compiler                                       |
|  Device  Tensor  Model  Stream  MemoryPool  PowerMonitor          |
|  QUAD IR (.qir)  QBin (.qbin)  ONNX frontend                      |
+-------------------------------------------------------------------+
|  Layer 1: Hardware                                                 |
|  CPU (Oryon/Kryo)   GPU (Adreno)   NPU (Hexagon HTP/DSP)          |
+-------------------------------------------------------------------+
""")

    h2(doc, "2.2  Mock vs. real adapter")
    p(doc, "QUAD ships in mock mode by default — every MCP tool returns "
           "deterministic synthetic data. This is what runs in CI and on a "
           "developer laptop without QAIRT. To switch to real mode you set "
           "the adapter to 'real' AND have a reachable QAIRT SDK.")
    table(doc, ["Mode", "Default", "What runs", "When to use"], [
        ["mock", "yes", "Deterministic Python, no subprocess", "CI, dev without SDK"],
        ["real", "no",  "Subprocess to qairt-converter / snpe-net-run / qairt-quantizer", "Production, QA, model bring-up"],
    ])

    # ─── 3. Hardware requirements ─────────────────────────────────────────
    h1(doc, "3. Hardware Requirements")
    h2(doc, "3.1  Recommended (real-mode end-to-end)")
    bullets(doc, [
        "Snapdragon X / X Elite / X2 Elite Copilot+ PC (Dell Latitude 7455, Lenovo ThinkPad T14s Gen 6, Microsoft Surface Pro 11/Laptop 7, HP OmniBook X, Samsung Galaxy Book4 Edge)",
        "16 GB+ RAM, 50 GB+ free disk (QAIRT alone is 1.7 GB compressed / 4.4 GB extracted)",
        "Windows 11 Pro (ARM64). PowerShell 7+ recommended; 5.1 also works.",
        "Stable internet for the one-time QAIRT download from the Qualcomm developer portal.",
    ])

    h2(doc, "3.2  Acceptable (mock-mode dev / CI)")
    bullets(doc, [
        "Any x86_64 Windows / macOS / Linux with Python 3.10+",
        "8 GB RAM, 5 GB free disk",
        "No SDK, no internet — works fully offline once installed",
    ])

    # ─── 4. Installation ──────────────────────────────────────────────────
    h1(doc, "4. Installation")
    h2(doc, "4.1  Windows ARM64 — the recommended path")
    p(doc, "On a fresh Snapdragon X-series Copilot+ PC, in PowerShell:")
    code(doc, """
git clone https://github.com/pkr465/QUAD.git
cd QUAD
.\\bootstrap.ps1 -QairtArchive C:\\Users\\<you>\\Downloads\\v2.46.0.260424.zip
""")
    p(doc, "What bootstrap.ps1 does:")
    numbered(doc, [
        "Installs Git Bash via winget if it isn't present (idempotent — re-runs are safe).",
        "Forwards to install.sh (POSIX) for the cross-platform install steps.",
        "Creates .venv\\, runs `pip install -e .[dev,real]`.",
        "Calls sdk_manager.install_archive on the QAIRT zip → unpacks to .\\sdks\\qairt-2.46.0.260424\\.",
        "Sets QAIRT_SDK_ROOT in the current session and writes activate.sh / activate.ps1 for later sessions.",
        "Runs `pytest -q` for verification.",
    ])

    callout(doc, "Re-running bootstrap.ps1 is always safe. It detects an existing "
                 "venv / SDK and skips work that's already done.", "tip")

    h2(doc, "4.2  macOS / Linux / WSL")
    code(doc, """
git clone https://github.com/pkr465/QUAD.git && cd QUAD
./install.sh --qairt-archive ~/Downloads/v2.46.0.260424.zip
source ./activate.sh && quad mode    # → 'real-mode: READY'
""")

    h2(doc, "4.3  Mock-only install (no SDK, no GPU/NPU access)")
    code(doc, """
./install.sh --mock-only    # macOS / Linux
.\\bootstrap.ps1 -MockOnly  # Windows PowerShell
""")
    p(doc, "Use this on CI runners and on developer laptops that don't need to "
           "talk to real hardware. The full test suite passes in this configuration "
           "— run `pytest -q` after install for the live count.")

    h2(doc, "4.4  CI / unattended install")
    p(doc, "Set QAIRT_DOWNLOAD_URL and QAIRT_DOWNLOAD_TOKEN to a pre-authorised "
           "internal mirror (the Qualcomm portal requires interactive EULA "
           "acceptance — there is no anonymous direct-download URL):")
    code(doc, """
export QAIRT_DOWNLOAD_URL=https://internal-mirror/qairt-2.46.0.260424.zip
export QAIRT_DOWNLOAD_TOKEN=Bearer-xxxxx
./install.sh    # Picks up the env vars automatically
""")

    h2(doc, "4.5  Verify")
    code(doc, """
quad mode                   # → adapter mode + real-mode readiness
quad sdk status             # → which SDK was discovered + version + bin dir
quad doctor                 # → 16 environment checks
quad doctor --real-mode     # → strict pre-flight; fails non-zero on any SDK issue
""")
    p(doc, "Expected output for a fully-real-mode install on a Snapdragon X box:")
    code(doc, """
adapter_mode:    real
sdk:             qairt 2.46.0.260424  (project:./sdks)
real-mode:       READY
  reason:        Real mode active. SDK root: ./sdks/qairt-2.46.0.260424
""")

    # ─── 5. SDK setup ─────────────────────────────────────────────────────
    h1(doc, "5. QAIRT SDK Setup")
    h2(doc, "5.1  Why a separate SDK?")
    p(doc, "The Qualcomm AI Engine Direct SDK (QAIRT) is the package that ships "
           "qairt-converter, qairt-quantizer, snpe-net-run, qnn-platform-validator "
           "and the libraries they depend on. QUAD's real adapter is a thin "
           "Python wrapper that shells out to those binaries; without QAIRT the "
           "real adapter has nothing to call.")

    h2(doc, "5.2  Downloading the SDK")
    numbered(doc, [
        "Sign in at https://www.qualcomm.com/developer/software/qualcomm-ai-engine-direct-sdk with a Qualcomm developer account.",
        "Accept the per-version EULA and download the zip (~1.7 GB).",
        "Save it anywhere — typically ~/Downloads/v2.46.0.260424.zip.",
    ])
    callout(doc, "The portal is a JS-rendered single-page app behind a login. "
                 "There is no anonymous direct-download URL. For unattended "
                 "installs, mirror the archive internally and set "
                 "QAIRT_DOWNLOAD_URL + QAIRT_DOWNLOAD_TOKEN.", "warning")

    h2(doc, "5.3  Installing")
    code(doc, """
quad sdk install ~/Downloads/v2.46.0.260424.zip
""")
    p(doc, "What this does:")
    numbered(doc, [
        "Validates the archive (zip-slip protection, format detection).",
        "Extracts to ./sdks/qairt-2.46.0.260424/ (gitignored).",
        "Auto-hoists nested wrapper directories so the bin/ tree is one level deep.",
        "Detects flavor (qairt vs snpe) and version from the install contents.",
        "Picks the per-arch bin/ subdir best suited to the host (Sprint 1 fix: ARM64 hosts get aarch64-windows-msvc, not the alphabetical-default x86_64).",
        "Persists discovery state to .quad/sdk.json.",
    ])

    h2(doc, "5.4  How discovery works")
    p(doc, "On every server start, sdk_manager.startup_resolve_and_log() searches in this order:")
    numbered(doc, [
        "QAIRT_SDK_ROOT / QNN_SDK_ROOT / SNPE_ROOT environment variables",
        "server.qairt_sdk_root from quad.toml",
        "./sdks/* (project-local)",
        "~/.quad/sdks/* (per-user)",
        "Vendor defaults: C:\\Qualcomm\\AIStack\\QAIRT\\*, /opt/qcom/aistack/qairt/*, /opt/qairt/*",
    ])
    p(doc, "First match wins. If nothing is found, the server runs in mock "
           "mode and prints the missing-SDK guidance with the exact URL the "
           "operator should visit.")

    h2(doc, "5.5  Multiple per-arch bin layouts")
    p(doc, "QAIRT 2.x splits its tools across per-architecture subdirectories:")
    table(doc, ["Subdir", "What's there", "When to use"], [
        ["aarch64-windows-msvc/", "Native ARM64 runtime: qnn-net-run.exe, snpe-net-run.exe, qnn-platform-validator.exe, snpe-diagview.exe", "ARM64 Windows hosts — fastest"],
        ["x86_64-windows-msvc/", "Full kit: converters AND runtime tools", "x86_64 Windows; ARM64 Windows via Prism emulation when a tool is missing from aarch64"],
        ["arm64x-windows-msvc/", "Mixed ARM64x: converters only", "ARM64 Windows when you want native converter execution"],
        ["x86_64-linux-clang/", "Linux x86_64 full kit", "Linux"],
        ["aarch64-ubuntu-gcc9.4/", "Linux ARM64 runtime", "Linux ARM64 hosts"],
    ])
    callout(doc, "QUAD's apply_to_environment() prepends every per-arch bin "
                 "subdir to PATH so a tool present only in aarch64-windows-msvc "
                 "is reachable even when the primary bin_dir points at "
                 "arm64x-windows-msvc.", "note")

    # ─── 6. Configuration ─────────────────────────────────────────────────
    h1(doc, "6. Configuration")
    h2(doc, "6.1  quad.toml (project-level)")
    code(doc, """
[server]
adapter_mode = "mock"          # "mock" or "real" (env var QUAD_ADAPTER_MODE overrides)
log_level = "info"
model_output_dir = "./output"

[adapters.qnn]
sdk_path = ""                  # Auto-discovered by sdk_manager — leave blank

[adapters.snpe]
sdk_path = ""                  # Same

[adapters.ai_hub]
api_key_env = "QAI_HUB_API_KEY"

[platforms.linux]
ssh_host = "arduino-uno-q.local"
ssh_user = "root"
""")

    h2(doc, "6.2  Environment variables")
    table(doc, ["Variable", "Purpose", "Default"], [
        ["QAIRT_SDK_ROOT", "Primary SDK install path", "(auto-discovered)"],
        ["QNN_SDK_ROOT", "Alternative — same effect", "(auto-discovered)"],
        ["SNPE_ROOT", "Legacy alternative", "(auto-discovered)"],
        ["QUAD_ADAPTER_MODE", "mock | real (overrides quad.toml)", "mock"],
        ["QUAD_STRICT_REAL", "1 = factory raises instead of falling back to mock", "0"],
        ["QUAD_SERVE_RUNTIME", "ModelServer runtime (mock | qairt) for `quad serve`", "(auto from QUAD_ADAPTER_MODE)"],
        ["QUAD_AIMET_BACKEND", "aimet_torch | aimet_onnx | qairt_quantizer | mock", "auto"],
        ["QUAD_PLACEHOLDER_BACKEND", "1 = compiler emits placeholder bytes (testing only)", "0"],
        ["QUAD_COMPILE_CACHE_DIR", "Override for compile cache location", "./.quad/compile_cache"],
        ["QAI_HUB_API_KEY", "Qualcomm AI Hub auth token", "(unset)"],
        ["ANDROID_SERIAL", "ADB device serial (Phase 3 platforms)", "(unset)"],
    ])

    # ─── 7. Running the MCP server ────────────────────────────────────────
    h1(doc, "7. Running the MCP Server")
    h2(doc, "7.1  Direct module entry")
    code(doc, """
# bash / Git Bash
./launch.sh
./launch.sh --real --verbose
./launch.sh --sse                          # SSE transport instead of stdio

# PowerShell
python -m quad.mcp.server
$env:QUAD_ADAPTER_MODE = "real"; python -m quad.mcp.server
""")
    callout(doc, "The legacy module path `python -m quad.server` still works "
                 "(backward-compat shim re-exports the FastMCP app). Prefer "
                 "`quad.mcp.server` in new code — that's the supported layout "
                 "after the quad/core / quad/mcp / quad/client split.", "note")

    h2(doc, "7.2  Claude Code auto-detection")
    p(doc, "install.sh writes a .claude/settings.json that Claude Code "
           "auto-detects when you open the repo:")
    code(doc, """
{
  "permissions": {
    "allow": [
      "mcp__quad__hardware_detect",
      "mcp__quad__convert_model",
      "mcp__quad__profile_workload",
      "mcp__quad__orchestrate_workload",
      "mcp__quad__generate_code"
    ]
  },
  "mcpServers": {
    "quad": {
      "command": "python",
      "args": ["-m", "quad.mcp.server"],
      "cwd": "${workspaceFolder}",
      "env": {
        "QUAD_ADAPTER_MODE": "mock"
      }
    }
  }
}
""")
    p(doc, "Open Claude Code and ask things like:")
    bullets(doc, [
        '"Detect the hardware on this machine."',
        '"Convert mobilenetv2.onnx to QNN INT8 and profile it on the NPU."',
        '"Generate Windows C++ inference code for this model."',
    ])

    # ─── 8. HTTP inference server ─────────────────────────────────────────
    h1(doc, "8. The Inference HTTP Server")
    p(doc, "`quad serve` starts a FastAPI app over uvicorn. As of Sprint 2 "
           "the server dispatches real inference through QAIRTAdapter when "
           "the model has a .dlc / .bin extension and runtime is qairt; "
           "otherwise it falls back to deterministic mock outputs.")

    h2(doc, "8.1  Starting the server")
    code(doc, """
# Mock-mode (default — returns shape-preserving random outputs)
quad serve mobilenetv2.dlc --name mnet --device npu --port 8080

# Real-mode — dispatch to snpe-net-run / qnn-net-run on the .dlc/.bin
QUAD_SERVE_RUNTIME=qairt quad serve mobilenetv2.dlc \\
    --name mnet --device npu --port 8080
""")
    p(doc, "The server prints which runtime it picked on startup. /health "
           "exposes the same label so monitoring can alert if a deploy lands "
           "with the wrong mode.")

    h2(doc, "8.2  Endpoints")
    table(doc, ["Method + path", "Purpose", "Returns"], [
        ["POST /infer", "Single inference", "{model_name, request_id, latency_ms, outputs}"],
        ["POST /infer/batch", "Batched inference (list of input dicts)", "list[InferResponse]"],
        ["GET /health", "Liveness + runtime tag", "{status, models_loaded, uptime_s, runtime}"],
        ["GET /metrics", "Aggregate stats", "{total_requests, avg_latency_ms, p99_latency_ms, throughput_rps, power_mw}"],
        ["GET /models", "List loaded models", "list[{name, path, device, version, num_inferences}]"],
        ["POST /models/{name}/load", "Load a model dynamically", "{status, model}"],
        ["DELETE /models/{name}", "Unload", "204"],
    ])

    h2(doc, "8.3  Example: posting an inference request")
    code(doc, """
import base64, numpy as np, httpx

x = np.random.rand(1, 3, 224, 224).astype(np.float32)
body = {
    "model_name": "mnet",
    "inputs": {
        "input": {
            "shape": list(x.shape),
            "dtype": str(x.dtype),
            "data_b64": base64.b64encode(x.tobytes()).decode(),
        },
    },
}
resp = httpx.post("http://localhost:8080/infer", json=body, timeout=30).json()
print(resp["latency_ms"], resp["outputs"]["output"]["shape"])
""")

    h2(doc, "8.4  Real-mode behaviour")
    p(doc, "Per Sprint 2: if the runtime is qairt and the model is .dlc/.bin, "
           "ModelServer delegates to QAIRTAdapter.execute_inference (which "
           "shells to snpe-net-run, marshals real input/output .raw files, and "
           "reads back numpy arrays). On subprocess failure the server falls "
           "back to mock and logs the failure — the HTTP layer never 500s "
           "on a single broken model while others work.")

    # ─── 9. Compiling ─────────────────────────────────────────────────────
    h1(doc, "9. Compiling Models")
    p(doc, "`quad compile` (and the underlying `quad.compiler.compile_model`) "
           "produces a .qbin fat binary containing IR plus per-target compiled "
           "binaries. Sprint 3 wired the real backend, so today the auto path "
           "shells to qairt-converter when an SDK is present and produces real "
           ".dlc bytes — not a placeholder.")

    h2(doc, "9.1  Backends")
    table(doc, ["Backend", "When picked", "What runs"], [
        ["auto", "Default — picks qairt if SDK is present, else stub", "Best path available"],
        ["qairt", "Explicit", "qairt-converter (+ qairt-quantizer for INT8/INT4)"],
        ["stub", "Explicit, or auto without SDK", "Raises BackendNotImplementedError (honest)"],
    ])

    h2(doc, "9.2  Examples")
    code(doc, """
# Compile mobilenetv2.onnx for the default targets, real backend (auto)
quad compile mobilenetv2.onnx --output mnet.qbin

# Force the real path explicitly with INT8 quantization
quad compile mobilenetv2.onnx --backend qairt --quantization int8

# Coverage-only run — no backend, just IR + per-target op-coverage report
quad compile resnet50.onnx --coverage-only
""")
    callout(doc, "compile_model caches by (sha256-of-source, target_sdk, "
                 "quantization). Re-compiling the same content is a single "
                 "dict lookup. Override the cache with use_cache=False or "
                 "$QUAD_COMPILE_CACHE_DIR=/tmp/empty for a clean rebuild.", "tip")

    # ─── 10. Quantization ─────────────────────────────────────────────────
    h1(doc, "10. Quantization")
    p(doc, "Sprint 4 added a third quantization backend, qairt_quantizer, "
           "that uses the SDK-bundled qairt-quantizer for real INT8/INT4. "
           "AIMET hooks remain in place but their wheels aren't on PyPI; "
           "qairt_quantizer is the credible default today.")

    h2(doc, "10.1  Backend selection priority (auto)")
    numbered(doc, [
        "aimet_torch (best — full PyTorch PTQ + QAT) — only if Qualcomm wheels are installed",
        "aimet_onnx (good — ONNX-only PTQ) — only if Qualcomm wheels are installed",
        "qairt_quantizer (decent — SDK-bundled PTQ, INT8/INT4)",
        "mock (deterministic, for tests)",
    ])

    h2(doc, "10.2  CLI examples")
    code(doc, """
# INT8 quantization with qairt-quantizer (default backend)
quad compile mobilenetv2.onnx --quantization int8

# INT4 with per-row block quantization (block size 32)
quad compile mobilenetv2.onnx --quantization int4
""")

    h2(doc, "10.3  Programmatic API")
    code(doc, """
from quad.adapters.aimet_adapter import AIMETAdapter, QuantizationConfig

adapter = AIMETAdapter(backend="qairt_quantizer", strict=True)
result = adapter.quantize(
    "model.dlc",
    output_path="model_int8.dlc",
    config=QuantizationConfig(bitwidth=8, scheme="symmetric_per_channel"),
    calibration="./calibration_data/",   # dir of .npy / .raw files
)
print(result.bitwidth, result.weight_size_compression, result.notes)
""")
    callout(doc, "accuracy_drop_estimate_pct on the result is a heuristic "
                 "based on Qualcomm public benchmarks. Production accuracy "
                 "claims must be measured on a real eval set.", "warning")

    # ─── 11. Profiling ────────────────────────────────────────────────────
    h1(doc, "11. Profiling")
    h2(doc, "11.1  Profiling levels")
    table(doc, ["Level", "What it measures", "Time"], [
        ["basic", "Total inference time", "Seconds"],
        ["detailed", "Per-layer timings via snpe-net-run --profiling_level detailed", "Tens of seconds"],
        ["linting", "HTP per-op cycle counts via snpe-net-run --profiling_level linting", "Tens of seconds"],
        ["qhas", "Full QHAS workflow (graph-prepare → net-run → qnn-profile-viewer chrometrace)", "Minutes"],
    ])

    h2(doc, "11.2  Examples")
    code(doc, """
quad profile mobilenetv2.dlc --profiling_level detailed
quad profile mobilenetv2.dlc --profiling_level linting    # bottleneck call-outs
quad profile mobilenetv2.dlc --profiling_level qhas        # chrometrace artifact
""")

    h2(doc, "11.3  Honest measurement reporting (Sprint 1)")
    p(doc, "ProfilingReport now carries a measurement_notes dict that tells "
           "callers whether each metric is measured, estimated, or not "
           "measured. snpe-net-run gives latency and per-op cycles; power and "
           "memory are NOT measured (no QPM3 integration yet) — the real "
           "adapter returns 0.0 for those, with a clear not_measured tag, "
           "instead of the historical fictional 2000 mW / 50 MB constants.")

    # ─── 12. Hardware detection ───────────────────────────────────────────
    h1(doc, "12. Hardware Detection")
    p(doc, "`quad detect` runs a real OS-level probe (PowerShell on Windows, "
           "/proc on Linux, sysctl on macOS, adb on Android when ANDROID_SERIAL "
           "is set) and reports actual CPU/GPU/NPU. As of Sprint 1, detect_hardware "
           "also parses qnn-platform-validator output for the supported runtimes, "
           "chipset, and Hexagon NPU details.")
    code(doc, """
quad detect              # cached
quad detect --refresh    # re-probe; bypasses the cache
""")

    # ─── 13. Code generation ──────────────────────────────────────────────
    h1(doc, "13. Code Generation")
    p(doc, "`quad generate` emits platform-specific inference code for QNN "
           ".so / .bin / TFLite-delegate pipelines. The QNN C++ templates have "
           "real init / load / execute / cleanup (no placeholder TODOs) and "
           "the validator will refuse to ship templates with TODO markers in "
           "function bodies.")
    code(doc, """
quad generate windows cpp mnet.dlc --output ./windows_qnn_app/
""")

    # ─── 14. Health, metrics, monitoring ──────────────────────────────────
    h1(doc, "14. Health, Metrics, Monitoring")
    p(doc, "When `quad serve` is running, scrape /metrics for per-deployment stats:")
    code(doc, """
curl http://localhost:8080/metrics
""")
    p(doc, "A typical Prometheus exporter / Grafana dashboard would track:")
    bullets(doc, [
        "avg_latency_ms (per model)",
        "p99_latency_ms",
        "throughput_rps",
        "models_loaded",
        "real-mode fallback events (in logs — see 'qairt_infer_failed_falling_back_to_mock')",
    ])

    # ─── 15. Troubleshooting ──────────────────────────────────────────────
    h1(doc, "15. Troubleshooting")
    table(doc, ["Symptom", "Likely cause", "Fix"], [
        ["`quad mode` says NOT READY", "QAIRT not discovered or env var stale", "Run `quad doctor --real-mode` — it prints the exact missing piece."],
        ["bootstrap.ps1 won't run (execution policy)", "PowerShell execution policy", "Use `bootstrap.bat`, or `powershell -ExecutionPolicy Bypass -File .\\bootstrap.ps1`."],
        ["`bash: command not found` on Windows", "Git Bash isn't installed", "Run bootstrap.ps1 first — it installs Git for Windows via winget."],
        ["./launch.sh not found in PowerShell", "Bash-style path", "Use `python -m quad.mcp.server` instead."],
        ["SDK installed but not detected", "SDK directory missing per-arch bin", "`quad sdk discover` shows what was scanned. Verify bin/<arch>/qairt-converter (or snpe-net-run) exists."],
        ["qnn-platform-validator missing on x86_64-windows-msvc", "QAIRT 2.46+ ships it only in aarch64-windows-msvc/", "Sprint-1 _find_tool walks every per-arch bin subdir; ensure apply_to_environment ran (it does on server start)."],
        ["Real-mode infer returns mock outputs", "Model extension not .dlc/.bin", "Sprint 2: real-mode dispatch is gated on the runnable extension. Compile the model first."],
        ["Real-mode infer falls back silently", "QAIRTAdapter subprocess failed", "Check `qairt_infer_failed_falling_back_to_mock` log entries; common causes: missing co-located libs, bad input shape."],
        ["Compile keeps re-running from scratch", "Cache disabled or content changed", "Confirm $QUAD_COMPILE_CACHE_DIR isn't pointing at a tmp-on-tmpfs that's wiped."],
        ["`pip install -e .[real]` fails on Windows ARM64", "Some real extras (psutil, paramiko) need wheels", "Use Python 3.12 (best wheel coverage); fall back to .[dev] only and add the real deps individually."],
    ])

    # ─── 16. Maintenance & upgrades ───────────────────────────────────────
    h1(doc, "16. Maintenance & Upgrades")
    h2(doc, "16.1  Bumping QAIRT")
    code(doc, """
# Download the new archive, then:
quad sdk install ~/Downloads/v2.47.0.260524.zip
# Restart the server. apply_to_environment refreshes QAIRT_SDK_ROOT.
""")
    h2(doc, "16.2  Upgrading QUAD itself")
    code(doc, """
git pull origin main
./install.sh        # idempotent — re-runs only what changed
pytest -q           # full suite — flag any new regressions
""")
    h2(doc, "16.3  Clearing the compile cache")
    code(doc, """
python -c "from quad.compiler.qairt_backend import clear_cache; print(clear_cache(), 'entries removed')"
""")

    # ─── 17. CI/CD ────────────────────────────────────────────────────────
    h1(doc, "17. CI/CD")
    h2(doc, "17.1  Three workflows")
    table(doc, ["Workflow", "When", "Purpose"], [
        ["ci.yml", "Push + PR to main", "Mock-mode unit tests on Ubuntu/Windows × Py 3.10–3.12, lint, type-check, package dry-run"],
        ["release.yml", "Tag push v*.*.* or manual", "Build wheel + sdist, smoke-install, publish to TestPyPI (manual) or PyPI (tag)"],
        ["real-hw.yml", "Manual + nightly + label `needs:real-hw`", "Self-hosted Snapdragon X runner: e2e + real-mode detect_hardware"],
    ])

    h2(doc, "17.2  Self-hosted runner")
    p(doc, "See docs/REAL_HARDWARE_CI.md for the full runbook. In short:")
    numbered(doc, [
        "Pre-stage the QAIRT zip on the runner (the workflow won't re-download a 1.7 GB archive each run).",
        "Register the runner with labels [self-hosted, snapdragon-x, windows].",
        "Set repo variable QAIRT_TEST_ARCHIVE = full path to the staged zip.",
        "Label-gate fork PRs (the workflow already requires `needs:real-hw`).",
    ])

    # ─── 18. CLI reference ────────────────────────────────────────────────
    h1(doc, "18. Appendix — CLI Reference")
    table(doc, ["Command", "What it does"], [
        ["quad mode", "Show adapter mode + real-mode readiness"],
        ["quad mode --set real", "Print 'export QUAD_ADAPTER_MODE=real' for shell-eval"],
        ["quad sdk status", "Show the active SDK + version + bin dir"],
        ["quad sdk discover", "Scan all standard locations, list every SDK"],
        ["quad sdk install <archive>", "Unpack a downloaded archive into ./sdks/"],
        ["quad doctor", "16 environment checks"],
        ["quad doctor --real-mode", "Strict pre-flight; exits non-zero on any SDK issue"],
        ["quad quickstart", "Interactive zero-to-inference wizard"],
        ["quad benchmark", "Standard benchmark suite (MobileNetV2 / ResNet / YOLOv8n)"],
        ["quad detect", "Real OS-level CPU / GPU / NPU + RAM probe"],
        ["quad detect --refresh", "Bypass cache and re-probe"],
        ["quad compile <model>", "Compile ONNX → .qbin (frontend real, backend qairt or stub)"],
        ["quad profile <model>", "Run platform profiler"],
        ["quad serve <model>", "Start FastAPI inference server"],
        ["quad configure", "Interactive SDK / target-device / API-key wizard"],
        ["quad-server", "FastMCP server entry point (long-running)"],
        ["quad-client", "Lightweight client provisioner (in QUAD-Client)"],
    ])

    # ─── 19. IoT Device Support ───────────────────────────────────────────
    h1(doc, "19. IoT Device Support")
    p(doc, "QUAD's adapter pattern extends from AI PC and Mobile to the full "
           "Qualcomm IoT SoC line. The full dependency catalogue (111 components "
           "across 14 categories) is published as a workbook at "
           "docs/IOT_DEPENDENCIES.xlsx. This section is a quick orientation; the "
           "workbook is authoritative.")

    h2(doc, "19.1  Target hardware")
    table(doc, ["Class", "SoC / Board", "Where it fits"], [
        ["Edge gateway",         "QCS6490 (RB3 Gen 2), QCS8250 (RB5)",  "Linux-class IoT, 12+ TOPS NPU"],
        ["High-perf IoT / AI",   "QCS8550",                              "Industrial AI, smart camera"],
        ["Camera / vision IoT",  "QCS610 / QCS605",                      "Smart camera, AI vision"],
        ["Entry IoT",            "QCM2290 / QCS2290",                    "Wearables, smart speakers"],
        ["Cellular IoT",         "Snapdragon X75 / 9205 modem",          "5G / NB-IoT / LTE-M uplink"],
    ])

    h2(doc, "19.2  Software dependency layers")
    table(doc, ["Layer", "What QUAD needs"], [
        ["OS / Firmware",        "Yocto meta-qcom, Qualcomm Linux, Zephyr / FreeRTOS, U-Boot, TF-A, OP-TEE"],
        ["Connectivity",         "BlueZ, OpenThread, Matter 1.4, hostapd, ModemManager + libqmi/libmbim"],
        ["IoT protocols",        "MQTT (Mosquitto / paho), CoAP (libcoap / aiocoap), LwM2M (Anjay), OPC-UA"],
        ["Cloud + OTA",          "AWS IoT Device SDK v2, azure-iot-device, Greengrass, IoT Edge, Mender, RAUC, SWUpdate"],
        ["Security",             "Qualcomm QTEE / SPU, OP-TEE, mbedTLS, OpenSSL, TF-M, PKCS#11, Matter Attestation"],
        ["Edge AI runtime",      "QNN / QAIRT 2.x, SNPE 2.x, Hexagon SDK 5.x, AIMET, Qualcomm AI Hub"],
        ["Sensors / HAL",        "libgpiod, smbus2, spidev, pyserial, python-can, pymodbus, bleak"],
        ["Telemetry",            "OpenTelemetry, Prometheus client_python, Fluent Bit"],
    ])

    callout(doc, "Open docs/IOT_DEPENDENCIES.xlsx for the full per-component "
                 "catalogue with priority (P1 / P2 / P3), license, target version, "
                 "and source URL. Filter by Category to scope a specific sub-stack.",
                 "tip")

    # ─── 20. Real-mode measurement plumbing ──────────────────────────────
    h1(doc, "20. Real-mode measurement plumbing")
    p(doc, "QUAD's profile_workload tool now returns a measurement_notes "
           "block that tags every metric with its provenance. The adapter "
           "auto-detects optional precision profilers (QPM3, Snapdragon "
           "Profiler) at runtime and silently uses measured values when "
           "they're installed; absent them, fields fall back to clearly "
           "labelled host-side estimates.")

    h2(doc, "20.1  measurement_notes — provenance tags")
    table(doc, ["Field", "Source tag", "Meaning"], [
        ["latency",     "measured:snpe-diagview",            "Parsed from snpe-diagview CSV (highest fidelity)"],
        ["latency",     "measured:snpe-net-run",             "Parsed from snpe-net-run stdout"],
        ["latency",     "not_measured:parser_no_match",      "Stdout / CSV didn't match any pattern (e.g. .onnx passed instead of .dlc)"],
        ["layers",      "measured:snpe-diagview",            "Per-layer rows from diagview CSV"],
        ["layers",      "synthetic_composite:no_diagview_csv","Fallback single-layer placeholder; diagview wasn't invoked"],
        ["memory",      "measured:psutil_rss(N_samples)",    "psutil polled the snpe-net-run subprocess every 100 ms"],
        ["power",       "measured:qpm3(N_samples)",          "QPM3 capture; per-frame PMIC readings"],
        ["power",       "estimated:host_thermal_model",      "Coarse model from CPU/GPU/NPU% × X-Elite TDP profile"],
        ["utilization", "measured:psutil_cpu+sdptrace_gpu",  "psutil for CPU, sdptrace chrometrace for GPU"],
        ["utilization", "measured:psutil_cpu_percent",       "CPU only (sdptrace absent)"],
        ["qhas_chrometrace", "measured:qnn-profile-viewer",  "QHAS chrometrace JSON path"],
    ])

    h2(doc, "20.2  Optional precision profilers")
    p(doc, "Two host-side tools, both gated downloads from Qualcomm:")
    table(doc, ["Tool", "Adds", "Detection / Activation"], [
        ["QPM3 (Qualcomm Power Monitor 3)",
         "Measured per-frame power (vs estimated)",
         "PATH or QPM3_HOME env var. Adapter auto-invokes during snpe-net-run."],
        ["Snapdragon Profiler (sdptrace)",
         "Real GPU% utilisation from chrometrace",
         "PATH or SNAPDRAGON_PROFILER_HOME. Adapter auto-captures concurrent trace."],
    ])
    callout(doc, "Neither tool blocks anything. Install once; quad doctor "
                 "flips them from WARN to PASS; profile_workload responses "
                 "automatically include the measured values. No client-side "
                 "changes needed.", "tip")

    h2(doc, "20.3  QAIRT-adapter stdout parsers")
    p(doc, "src/quad/adapters/parsers.py contains four pure-function parsers "
           "wired into the adapter:")
    bullets(doc, [
        "parse_snpe_net_run_stdout — total/forward inference time, runtime banner, error tags",
        "parse_snpe_diagview_csv — Total Inference Time, Forward Propagate, per-section init metrics",
        "parse_snpe_diagview_layers — per-layer rows from the 'Model Layer Times' section",
        "parse_qnn_platform_validator — per-backend block parser (DSP / GPU / CPU support state)",
        "parse_qairt_converter_stdout — supported_ops_pct, unsupported_ops list, warnings, errors",
    ])

    h2(doc, "20.4  Model registry — `quad models`")
    p(doc, "Production ONNX provisioning lives in src/quad/model_registry/. "
           "Models are listed in registry.yaml with a name, plan, and either a "
           "url (auto-downloadable) or a path_env_var (user-supplied for "
           "gated / large weights).")
    code(doc, """
quad models list                  # show every entry with cache state
quad models fetch mobilenetv2     # download + verify SHA-256
quad models path llama3_8b_prefill # resolve $LLAMA3_8B_PREFILL_ONNX
quad models verify <name>         # re-check SHA-256 of cached file
""")
    p(doc, "Adding a model for a new plan is a single yaml entry — no Python "
           "changes. The fetcher streams via httpx, atomic-renames into the "
           "cache (~/.quad/models/), and verifies SHA-256 when declared.")

    h2(doc, "20.5  Snapdragon X Elite environment caveat")
    p(doc, "On Snapdragon X Elite (Windows-on-ARM64), QAIRT 2.46's host "
           "Python tools (qairt-converter, qairt-quantizer) ship as "
           "windows-arm64ec/.pyd modules that depend on the full Visual "
           "Studio 2022 runtime, not just the VC++ redistributable. The "
           "RUNTIME path (snpe-net-run, qnn-platform-validator) works on a "
           "stock install — only model conversion is affected. To enable "
           "conversion on this box: `winget install "
           "Microsoft.VisualStudio.2022.Community` (or use BuildTools), or "
           "convert on a separate x86_64 host and copy the .dlc back. "
           "`quad doctor --real-mode` surfaces this explicitly under 'Python "
           "arch vs OS'.")

    # ─── End ──────────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "Reference Documents")
    bullets(doc, [
        "docs/PRODUCTION_READINESS_REVIEW_2026-05-08.md — the gap analysis that drove Sprints 1–5",
        "docs/REAL_HARDWARE.md — one-step real-mode enablement playbook",
        "docs/REAL_HARDWARE_CI.md — self-hosted runner setup",
        "docs/SAMPLE_APP_REPORT.md — Snapdragon X Elite measurements",
        "docs/IOT_DEPENDENCIES.xlsx — IoT device support catalogue (111 components)",
        "src/quad/adapters/parsers.py — QAIRT/SNPE stdout parsers (snpe-net-run, snpe-diagview, qnn-platform-validator, qairt-converter)",
        "src/quad/profiler/{qpm3,sdptrace,host_power,host_utilization,rss_sampler}.py — optional precision profilers + host-side measurement helpers",
        "src/quad/model_registry/ — production ONNX provisioning (registry.yaml + fetcher.py)",
        "tests/e2e/test_real_sdk_e2e.py — the canonical 7-phase e2e validation",
        "QUAD_Client_Guide.docx (companion) — the developer-facing client guide",
    ])

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Wrote {out_path}")


# ═══════════════════════════════════════════════════════════════════════════
# CLIENT GUIDE
# ═══════════════════════════════════════════════════════════════════════════


def build_client_guide(out_path: Path) -> None:
    doc = Document()
    _setup_styles(doc)

    title_page(
        doc,
        "QUAD Client Guide",
        "Developer playbook for using QUAD MCP server through Claude Code",
        "Application developers, ML engineers, anyone using Claude Code on Qualcomm hardware",
    )

    # ─── 1. What this is ──────────────────────────────────────────────────
    h1(doc, "1. What QUAD-Client Does")
    p(doc, "QUAD-Client is a 6 MB Python provisioner that wires Claude Code "
           "(or another MCP-compatible IDE) to a QUAD MCP server. It does NOT "
           "include the QAIRT SDK, the compiler, or the runtime — those live "
           "on the server. The client just generates the right "
           ".claude/settings.json and bundles 11 skill files.")

    h2(doc, "1.1  Three deployment topologies")
    table(doc, ["Topology", "Where the server runs", "Transport"], [
        ["Local", "Same machine as the IDE", "stdio-local"],
        ["Lab box over SSH", "Remote machine you ssh to", "stdio-ssh"],
        ["Hosted MCP service", "Service over HTTP/SSE", "sse-http"],
    ])
    callout(doc, "Pick the topology that matches your environment. The client "
                 "supports all three; the install commands differ only in the "
                 "--transport flag.", "tip")

    h2(doc, "1.2  Goal")
    p(doc, "From a clean machine you should be able to chat with Claude about "
           "your Qualcomm models — convert, profile, allocate, generate code "
           "— in well under five minutes. This guide walks you through.")

    # ─── 2. Quick start ───────────────────────────────────────────────────
    h1(doc, "2. Quick Start (30-second install)")
    h2(doc, "2.1  Option A — pip install (recommended once published)")
    code(doc, """
pip install quad-mcp-client
quad-client install        # interactive: prompts for transport + tests connection
""")
    h2(doc, "2.2  Option B — clone + run installer")
    code(doc, """
git clone https://github.com/pkr465/QUAD-Client.git
cd QUAD-Client
./install.sh                  # Linux / macOS / Git Bash on Windows
# or:
.\\bootstrap.ps1               # Windows PowerShell
""")
    h2(doc, "2.3  What gets created")
    table(doc, ["Path", "Purpose", "Size"], [
        ["./.claude/settings.json", "Tells Claude Code where the QUAD MCP server is", "~1 KB"],
        ["./.claude/skills/quad-*.md", "11 bundled skill files (quickstart, detect, convert, profile, …)", "~50 KB"],
        ["~/.local/lib/.../quad_mcp_client/", "Python package (typer + httpx, no heavy deps)", "~5 MB"],
    ])
    callout(doc, "Nothing else is installed. No fastmcp, numpy, or pydantic. "
                 "A subprocess test in CI enforces this — importing "
                 "quad_mcp_client.cli is verified to pull zero heavy modules.", "note")

    # ─── 3. Three transports ──────────────────────────────────────────────
    h1(doc, "3. The Three Transports")
    h2(doc, "3.1  stdio-local (server on the same machine)")
    p(doc, "Use when the QUAD server is installed locally — the most common "
           "case for a single-developer workstation.")
    code(doc, """
quad-client install --transport stdio-local
""")
    p(doc, "Generates a settings.json that tells Claude Code to launch the "
           "MCP server via stdio:")
    code(doc, """
{
  "mcpServers": {
    "quad": {
      "command": "python",
      "args": ["-m", "quad.mcp.server"],
      "cwd": "${workspaceFolder}",
      "env": { "QUAD_ADAPTER_MODE": "mock" }
    }
  }
}
""")
    p(doc, "Pre-flight check (Sprint 1 P0-8): the install command first runs "
           "`python -c 'import sys; assert sys.version_info >= (3,10); "
           "import quad.mcp.server'` to verify the right Python is on PATH "
           "and the QUAD package is importable. If either fails, the install "
           "aborts with a clear hint instead of writing a broken settings.json.")

    h2(doc, "3.2  stdio-ssh (server on a remote lab box)")
    p(doc, "Use when QUAD lives on a beefy lab machine and you're working "
           "from a thin laptop. Claude Code spawns ssh, and ssh runs the QUAD "
           "server on the remote — no port-forward needed.")
    code(doc, """
quad-client install --transport stdio-ssh \\
    --ssh-user alice --ssh-host lab.example.com
""")
    p(doc, "Optional flags:")
    table(doc, ["Flag", "Meaning"], [
        ["--ssh-port 22", "Override SSH port"],
        ["--ssh-key ~/.ssh/quad_key", "Use a specific identity file"],
        ["--server-command 'python -m quad.mcp.server'", "Override the remote command (rare)"],
    ])
    p(doc, "The probe runs in batch mode — no interactive password prompts "
           "— so misconfigured key auth fails fast with this kind of hint:")
    code(doc, """
✗ stdio-ssh probe failed: ssh exit 255: Permission denied (publickey).
  → SSH key auth to alice@lab.example.com failed. Add your public key to
    the server's ~/.ssh/authorized_keys, or specify a key with --ssh-key.
""")

    h2(doc, "3.3  sse-http (hosted MCP service)")
    p(doc, "Use when QUAD is exposed as a managed HTTP/SSE endpoint — e.g. "
           "an internal team service. Bearer-token auth supported.")
    code(doc, """
quad-client install --transport sse-http \\
    --sse-url https://mcp.example.com/sse \\
    --sse-auth-token-env QUAD_MCP_TOKEN
""")
    p(doc, "Sprint 1 hardening: the probe checks the response Content-Type "
           "in addition to status code. text/event-stream and "
           "application/json are accepted; text/html is rejected with a "
           "clear 'wrong URL' hint.")

    # ─── 4. Verifying the connection ──────────────────────────────────────
    h1(doc, "4. Verifying the Connection")
    code(doc, """
# Without writing settings — just probe
quad-client connect-test stdio-local
quad-client connect-test stdio-ssh --ssh-user alice --ssh-host lab.example.com
quad-client connect-test sse-http --sse-url https://mcp.example.com/sse

# Show what would be written to settings.json without touching the file
quad-client preview --transport sse-http --sse-url https://mcp.example.com/sse

# Show current install status
quad-client status
""")

    # ─── 5. The 5 MCP tools ───────────────────────────────────────────────
    h1(doc, "5. The Five MCP Tools")
    p(doc, "Once the client is installed and Claude Code is open, you have "
           "five tools available. You don't call them by name — Claude picks "
           "the right one based on your prompt. Here's what each does and "
           "the kind of phrasing that triggers it.")

    h2(doc, "5.1  hardware_detect")
    table(doc, ["Inputs", "Output"], [
        ["platform: windows | linux | android", "DeviceProfile JSON: chipset, CPU/GPU/NPU specs, available SDK"],
    ])
    p(doc, "Trigger phrases:")
    bullets(doc, [
        '"What hardware do I have?"',
        '"Is this Snapdragon?"',
        '"Do I have an NPU?"',
        '"Detect the chipset on this machine."',
    ])

    h2(doc, "5.2  convert_model")
    table(doc, ["Inputs", "Output"], [
        ["source_format: onnx|pytorch|tensorflow|tflite\nmodel_path\ntarget_sdk: qnn|snpe\nquantization: fp32|int8|int4", "Output path (.bin for QNN, .dlc for SNPE) + conversion notes + image-format guidance"],
    ])
    p(doc, "Trigger phrases:")
    bullets(doc, [
        '"Convert mobilenetv2.onnx to QNN."',
        '"Quantize this model to INT8."',
        '"Compile resnet.onnx for SNPE on Linux."',
    ])

    h2(doc, "5.3  profile_workload")
    table(doc, ["Inputs", "Output"], [
        ["model_path\nplatform\nruntime: cpu|gpu|npu|auto\nprofiling_level: basic|detailed|linting|qhas", "Latency stats (mean/p50/p95/p99), throughput, memory, per-layer or per-op profile"],
    ])
    p(doc, "Trigger phrases:")
    bullets(doc, [
        '"Profile mnet.dlc on the NPU."',
        '"Where are the bottlenecks?"',
        '"Run linting on this model."',
        '"Generate a QHAS chrometrace."',
    ])

    h2(doc, "5.4  orchestrate_workload")
    table(doc, ["Inputs", "Output"], [
        ["model_path\npower_mode: performance|balanced|efficiency", "Per-layer allocation map (CPU/GPU/NPU) + utilization estimates"],
    ])
    p(doc, "Trigger phrases:")
    bullets(doc, [
        '"How should I split this across CPU/GPU/NPU?"',
        '"Compare power modes for this model."',
        '"Optimise for battery life."',
    ])

    h2(doc, "5.5  generate_code")
    table(doc, ["Inputs", "Output"], [
        ["platform: windows|linux|android\nlanguage: cpp|python|kotlin|arduino_sketch\nsdk: qnn|snpe\nmodel_path", "Source files + build instructions"],
    ])
    p(doc, "Trigger phrases:")
    bullets(doc, [
        '"Generate Windows C++ inference code for mnet.dlc."',
        '"Give me an Android AAR for this model."',
        '"Emit Python inference code for Linux."',
    ])

    # ─── 6. The 11 skills ─────────────────────────────────────────────────
    h1(doc, "6. The 11 Bundled Skills")
    p(doc, "Each skill is a markdown file in .claude/skills/quad-*.md. Claude "
           "loads them as routing hints — when you say something that matches "
           "a skill's trigger phrases, Claude applies that skill's "
           "instructions when picking which tools to call and how to format "
           "the result.")
    table(doc, ["Skill", "Trigger phrases", "What it adds"], [
        ["quad-quickstart", '"get started", "set up QUAD"', "End-to-end walkthrough: detect → convert → profile → orchestrate → codegen"],
        ["quad-detect", '"what hardware", "is this Snapdragon"', "Real probe + Qualcomm-vs-other tip routing"],
        ["quad-convert", '"convert this model", "quantize"', "Calibration data + image-format guidance"],
        ["quad-profile", '"profile", "find bottlenecks", "QHAS"', "Picks the right level + bottleneck callouts"],
        ["quad-orchestrate", '"allocate across CPU/GPU/NPU", "compare power modes"', "3-mode comparison + fallback analysis"],
        ["quad-codegen", '"generate inference code", "C++ for this model"', "Platform/language/sdk picker + build commands"],
        ["quad-doctor", '"is QUAD set up", "diagnose"', "Diagnostic translation table — every check has an exact fix command"],
        ["quad-deploy", '"deploy to my phone", "push to Arduino"', "deploy.sh + remote profiling walkthrough"],
        ["quad-recommend", '"what\'s the best way", "INT4 or INT8"', "Synthesises model + target + use case into a categorised plan"],
        ["quad-serve", '"start an inference server", "expose as HTTP API"', "FastAPI server setup + curl + Python client snippets"],
        ["quad-configure", '"configure", "set API keys"', "SDK / target-device / API-key wizard"],
    ])

    # ─── 7. End-to-end workflows ──────────────────────────────────────────
    h1(doc, "7. End-to-End Workflows")
    h2(doc, "7.1  Workflow 1 — ONNX → running on NPU in 5 minutes")
    p(doc, "Open Claude Code in your repo. Try a session like this.")
    code(doc, """
You: I have mobilenetv2.onnx. Get it running on my NPU and tell me the
     fastest power mode.

Claude: I'll detect your hardware first, then convert and profile.
        [Calls hardware_detect → Snapdragon X Elite, Hexagon NPU]
        [Calls convert_model(onnx → QNN INT8) → mnet.dlc + notes]
        [Calls profile_workload(detailed) → 2.6 ms p50]
        [Calls orchestrate_workload(performance) → 95% on NPU]
        Conversion produced mnet.dlc (3.4 MB, ~4× compression vs FP32).
        Profile: 2.6 ms p50, 388 FPS, 95% NPU utilisation. Performance
        mode is the right choice here — balanced gives up 12% latency
        for ~40% power saving; efficiency mode falls back to CPU which
        is 8× slower.
""")

    h2(doc, "7.2  Workflow 2 — Profile a slow model and find bottlenecks")
    code(doc, """
You: My model is at 30 ms but I need under 10 ms. What's slow?

Claude: Running linting profile to get per-op cycles…
        [Calls profile_workload(linting)]
        Top three bottlenecks (cycle fractions):
          1. Conv2d_42  — 18% of total cycles, large kernel (7×7),
             SoftReLU activation forces fallback to CPU. Replace with
             3 stacked 3×3 + ReLU; expect ~12 ms drop.
          2. Concat_18  — 9%, awkward axis. Reorder upstream layers
             to align channels.
          3. LayerNorm  — 7%, single-op overhead. Fuse into the
             preceding MatMul if possible (HTP supports the fused op).
""")

    h2(doc, "7.3  Workflow 3 — Generate Windows C++ inference code")
    code(doc, """
You: Emit the Windows C++ that runs mnet.dlc end-to-end.

Claude: [Calls generate_code(platform=windows, language=cpp, sdk=qnn)]
        Wrote to ./generated/mnet_windows/:
          inference.cpp  — full QNN init/load/execute/cleanup
          inputs.cpp     — pre-processing (NHWC, BGR mean subtract)
          CMakeLists.txt — links libQnnHtp + libQnnSystem
        Build with:
          cmake -B build && cmake --build build --config Release
        Run:
          .\\build\\Release\\mnet_infer.exe inputs/test.png
""")

    h2(doc, "7.4  Workflow 4 — Compare CPU vs GPU vs NPU")
    code(doc, """
You: How do power and latency compare across runtimes for mnet?

Claude: [Calls profile_workload x3, runtime=cpu / gpu / npu]
        Runtime  | Latency  | Power (typ.) | Use case
        ---------|----------|--------------|----------------
        CPU      | 22.1 ms  |  3000 mW     | Universal fallback
        GPU      |  8.7 ms  |  4500 mW     | Compute-rich, mid-power
        NPU      |  2.6 ms  |  1500 mW     | Best perf/W — ship this
""")

    # ─── 8. Common issues ─────────────────────────────────────────────────
    h1(doc, "8. Common Issues")
    table(doc, ["Symptom", "Fix"], [
        ["Claude says 'no QUAD MCP server detected'", "`quad-client status` to see what's installed; re-run `quad-client install`."],
        ["stdio-local probe says 'python too old'", "Sprint 1 P0-8 catches this. Install Python 3.10+ and put it first on PATH, or pass --command <full-path-to-python>."],
        ["stdio-local probe says 'quad.mcp.server not importable'", "Activate the venv where you installed quad-agent before running `quad-client install`."],
        ["stdio-ssh fails with 'Permission denied (publickey)'", "Add your public key to the server's ~/.ssh/authorized_keys, or pass --ssh-key /path/to/key."],
        ["sse-http says 'HTTP 200 but unexpected content-type: text/html'", "Sprint 1 P0-9: the URL points at a static page, not the MCP endpoint. Check the path — typically /sse or /mcp."],
        ["Claude calls a tool but gets a fall-back mock response", "The server is in mock mode. Remote: ssh in, run `quad mode --set real`. Local: set `QUAD_ADAPTER_MODE=real` in the settings.json env block and reload."],
        ["A specific MCP tool isn't allowed", "Check .claude/settings.json `permissions.allow` includes mcp__quad__<tool_name>."],
    ])

    # ─── 9. CLI reference ─────────────────────────────────────────────────
    h1(doc, "9. CLI Reference")
    table(doc, ["Command", "What it does"], [
        ["quad-client install", "Provision: probe, generate settings.json, copy skills"],
        ["quad-client install --transport <kind>", "Skip the interactive prompt"],
        ["quad-client install --force", "Skip the connection probe (write anyway)"],
        ["quad-client install --skip-test", "Skip the probe, but keep prompts"],
        ["quad-client status", "What's installed in the current dir"],
        ["quad-client preview --transport <kind> ...", "Print settings.json without writing"],
        ["quad-client uninstall", "Remove bundled skills (settings.json kept)"],
        ["quad-client connect-test <transport> ...", "Standalone probe — no provisioning"],
    ])
    p(doc, "Run any subcommand with --help for the full flag list.")

    # ─── 10. Uninstall ────────────────────────────────────────────────────
    h1(doc, "10. Uninstalling")
    code(doc, """
quad-client uninstall          # removes the 11 bundled skills
rm .claude/settings.json       # only if you want to fully detach Claude Code
pip uninstall quad-mcp-client
""")

    # ─── 11. FAQ ──────────────────────────────────────────────────────────
    h1(doc, "11. FAQ")
    h3(doc, "Does this work with Cursor / Continue / Cline?")
    p(doc, "Today: Claude Code only. Adding a new IDE means implementing "
           "MCPClientProvisioner under src/quad_mcp_client/<client>/. The "
           "abstract base class is already in place — see "
           "docs/DEVELOPER_GUIDE.md.")

    h3(doc, "Is there a MCP-over-WebSocket transport?")
    p(doc, "Not yet. SSE-over-HTTPS covers most managed-service deploys. "
           "Open an issue if you have a strong WS use case.")

    h3(doc, "Why no fastmcp / numpy / pydantic in the client?")
    p(doc, "The client just emits a JSON file and runs a small probe. "
           "Claude Code is what actually speaks MCP. Keeping the client at "
           "~6 MB means we can ship it as a standalone wheel with a tiny "
           "dependency tree — typer + httpx, that's it.")

    h3(doc, "Can the server and client be different versions?")
    p(doc, "Yes. The client only depends on the five MCP tool names. The "
           "server is free to evolve internals as long as the tool surface is "
           "preserved. We test the cross-version compatibility nightly via "
           "real-hw.yml.")

    h3(doc, "How do I switch between mock and real mode?")
    p(doc, "Edit .claude/settings.json and set "
           '"env": {"QUAD_ADAPTER_MODE": "real"} — Claude Code will pick it up '
           "on the next reload. The server will fall back to mock if the SDK "
           "isn't reachable, with a tagged warning. Check `quad doctor "
           "--real-mode` on the server side to see why.")

    h3(doc, "What's in measurement_notes when I call profile_workload?")
    p(doc, "Every metric is tagged with its provenance so you can tell at a "
           "glance whether QUAD measured it or estimated it. Tags include "
           "measured:snpe-diagview, measured:psutil_rss(N_samples), "
           "measured:qpm3(N_samples), estimated:host_thermal_model, "
           "and not_measured:parser_no_match. On a server without QPM3 / "
           "Snapdragon Profiler installed, power reads "
           "estimated:host_thermal_model and utilisation is psutil-only "
           "(no GPU%). Installing those tools server-side flips the tags "
           "to measured without any client change.")

    h3(doc, "How do I provision real ONNX models?")
    p(doc, "Use the server's `quad models` registry. From the IDE side you "
           "can ask Claude:")
    code(doc, """
> Use QUAD to list every model in the registry and tell me which are cached
> Use QUAD to fetch mobilenetv2 from the registry, then run hardware_detect
> on Snapdragon X Elite and profile_workload on the cached model
""")
    p(doc, "For gated weights (Llama 3 8B), set the matching env var on the "
           "server side (LLAMA3_8B_PREFILL_ONNX, LLAMA3_8B_DECODE_ONNX) and "
           "the registry entry will resolve to that local path.")

    h3(doc, "Does QUAD-Client work with Qualcomm IoT boards (RB3 Gen 2 / RB5)?")
    p(doc, "Yes — QUAD-Client is transport-agnostic. Point it at a server running "
           "on the IoT board (or on a workstation that talks to it over SSH / "
           "ADB / serial). The five MCP tools work identically; the server-side "
           "adapters handle the SoC-specific bits. The full IoT dependency "
           "catalogue lives in the QUAD repo at docs/IOT_DEPENDENCIES.xlsx.")

    h1(doc, "Reference")
    bullets(doc, [
        "QUAD_Server_Guide.docx — the operator-facing companion",
        "docs/INSTALL.md — install detail per topology",
        "docs/TRANSPORTS.md — protocol details for each transport",
        "docs/TROUBLESHOOTING.md — full error catalogue",
        "docs/DEVELOPER_GUIDE.md — adding new IDE clients, contributing",
        "QUAD repo → docs/IOT_DEPENDENCIES.xlsx — IoT device support catalogue",
    ])

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Wrote {out_path}")


# ═══════════════════════════════════════════════════════════════════════════


if __name__ == "__main__":
    here = Path(__file__).resolve()
    repo_root = here.parent.parent  # docs/.. = QUAD root
    server_out = repo_root / "docs" / "QUAD_Server_Guide.docx"
    client_repo = repo_root.parent / "QUAD-Client"
    client_out = client_repo / "docs" / "QUAD_Client_Guide.docx"

    build_server_guide(server_out)
    if client_repo.is_dir():
        build_client_guide(client_out)
    else:
        # Fall back to writing the client guide alongside the server guide
        # so a single-repo checkout still produces both files.
        build_client_guide(server_out.with_name("QUAD_Client_Guide.docx"))
