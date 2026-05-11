"""Generate Current_Status_And_Roadmap.docx — a dated snapshot of where
QUAD is and what's pending for commercial launch.

Format mirrors the other QUAD design docs (generate_user_guides.py).
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent / "Current_Status_And_Roadmap.docx"
SNAPSHOT_DATE = date.today()


# ── Style helpers ─────────────────────────────────────────────────────────


def setup_styles(doc: Document) -> None:
    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(11)
    for h, sz in (("Heading 1", 22), ("Heading 2", 16), ("Heading 3", 13)):
        st = doc.styles[h]
        st.font.name = "Calibri"
        st.font.size = Pt(sz)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5C)


def shade(par, hex_color: str) -> None:
    p_pr = par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    p_pr.append(shd)


def callout(doc: Document, text: str, kind: str = "note") -> None:
    palette = {
        "note":    ("E8F0FE", "1A73E8", "Note"),
        "warning": ("FCE8E6", "C5221F", "Warning"),
        "tip":     ("E6F4EA", "188038", "Tip"),
    }
    bg, fg, label = palette[kind]
    par = doc.add_paragraph()
    r = par.add_run(f"  {label}: ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(fg)
    par.add_run(text)
    shade(par, bg)


def table(doc, headers, rows) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        t.style = "Light Grid Accent 1"
    except KeyError:
        pass
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            c = t.rows[r_idx + 1].cells[c_idx]
            c.text = ""
            r = c.paragraphs[0].add_run(val)
            r.font.size = Pt(10)


def h1(d, t): d.add_heading(t, level=1)
def h2(d, t): d.add_heading(t, level=2)
def h3(d, t): d.add_heading(t, level=3)
def p(d, t):
    par = d.add_paragraph()
    par.add_run(t)


def bullets(d, items):
    for it in items:
        try:
            par = d.add_paragraph(style="List Bullet")
        except KeyError:
            par = d.add_paragraph()
            par.add_run("•  ")
        par.add_run(it)


# ── Document content ──────────────────────────────────────────────────────


def title_page(doc):
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(140)
    r = p1.add_run("QUAD — Current Status & Roadmap")
    r.bold = True
    r.font.size = Pt(32)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5C)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Honest snapshot of supported features, validation gaps, and the path to commercial v1.0")
    r2.italic = True
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(80)
    r3 = p3.add_run(f"Snapshot date: {SNAPSHOT_DATE.isoformat()}    |    QUAD v0.4.0")
    r3.font.size = Pt(12)
    r3.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_page_break()


def section_sources_of_truth(doc):
    h1(doc, "Sources of truth")
    p(doc,
      "This snapshot is assembled from the heads of pkr465/QUAD and "
      "pkr465/QUAD-Client as of the date on the title page, plus the "
      "validation results captured by projects/real_hw/smoke.py on a "
      "Snapdragon X Elite reference box. Where this document and "
      "CLAUDE.md disagree, CLAUDE.md is authoritative — it updates "
      "every session.")
    table(doc, ["Source", "What it carries"], [
        ["QUAD/CLAUDE.md",                        "Phase status, Active Context, decisions log"],
        ["QUAD/docs/QUAD_Platform_Design_v2.docx","7-layer architecture, design rationale"],
        ["QUAD/docs/QUAD_Server_Guide.docx",      "Operator playbook (server-side); Section 20 covers real-mode measurement plumbing"],
        ["QUAD-Client/docs/QUAD_Client_Guide.docx","Developer playbook (IDE-side)"],
        ["projects/NPU_Implementation_Plans.docx","Six implementation plans (4 features + 2 audience walkthroughs)"],
        ["QUAD-Academy/",                         "Curriculum + labs + certification programme"],
        ["projects/real_hw/real_hw_report.json",  "Real-mode smoke output from the X Elite reference box"],
    ])


def section_supported(doc):
    h1(doc, "Currently supported")

    h2(doc, "Core MCP tool surface")
    table(doc, ["Tool", "Mock mode", "Real mode", "Notes"], [
        ["hardware_detect", "Yes", "Yes",
         "Real X Elite probe verified: chipset, NPU, GPU, RAM, runtimes, SDK version"],
        ["convert_model", "Yes", "Yes (wired)",
         "qairt-converter invokes end-to-end on X Elite (with VS 2022 + Py 3.10 + dlc_utils patch); --source_model_input_shape emitted; SHA-verified output"],
        ["profile_workload (detailed)", "Yes", "Yes",
         "snpe-net-run + snpe-diagview CSV parser; provenance-tagged latency / per-layer"],
        ["profile_workload (linting)", "Yes", "Yes",
         "HTP cycle counts, bottleneck analysis, optimisation hints"],
        ["profile_workload (qhas)", "Yes", "Yes",
         "qnn-profile-viewer chrometrace, GPU% from chrometrace events"],
        ["orchestrate_workload", "Yes", "Yes",
         "CPU/GPU/NPU allocation, power-mode-aware projections"],
        ["generate_code", "Yes", "Yes",
         "C++ / Python / Kotlin; multi-context QNN pipeline template"],
    ])

    h2(doc, "Measurement plumbing (provenance-tagged)")
    table(doc, ["Metric", "Source", "Status"], [
        ["Latency",
         "snpe-diagview CSV → fallback snpe-net-run stdout",
         "Wired"],
        ["Per-layer",
         "snpe-diagview CSV → fallback synthetic-composite",
         "Wired"],
        ["Memory RSS",
         "psutil async sampler (100 ms polling)",
         "Wired"],
        ["CPU%",
         "psutil.cpu_percent",
         "Wired"],
        ["NPU%",
         "Arithmetic from linting cycle counts / wall_time / Hexagon clock",
         "Wired"],
        ["GPU%",
         "Snapdragon Profiler chrometrace event sum",
         "Auto-activated when sdptrace installed"],
        ["Power (estimated)",
         "Host thermal model: CPU/GPU/NPU% × X Elite TDP profile",
         "Default fallback"],
        ["Power (measured)",
         "QPM3 per-frame PMIC reading",
         "Auto-activated when QPM3 installed"],
    ])

    h2(doc, "Platform / runtime support")
    table(doc, ["Item", "Status"], [
        ["Snapdragon X Elite (Windows ARM64)",
         "Runtime verified end-to-end; conversion needs VS 2022 + Py 3.10 (auto-installed by bootstrap.ps1)"],
        ["Snapdragon 8 Elite (Android)",      "Adapter wired (mock); real-hw untested"],
        ["Arduino UNO Q / QCS2210",          "Adapter wired (mock); real-hw untested"],
        ["Linux x86_64",                     "Mock + real-mode dispatch ready"],
        ["QAIRT 2.46+",                       "Verified"],
        ["SNPE 2.x",                         "Adapter present"],
        ["Hexagon SDK 5.x",                  "Kernel DSL → qbin lowering path exists (Phase E)"],
        ["AIMET INT8 / INT4",                "Adapter present; not exercised on real HW for INT4 LLMs"],
        ["Qualcomm AI Hub",                  "Adapter present (mock); cloud calls untested"],
    ])

    h2(doc, "Developer experience")
    table(doc, ["Item", "Status"], [
        ["quad CLI (configure / quickstart / doctor / mode / sdk / detect / compile / optimize / profile / serve / benchmark / models)",
         "Shipped"],
        ["quad doctor --real-mode (19 checks)",
         "Shipped"],
        ["quad models {list,fetch,path,verify} registry CLI",
         "Shipped"],
        ["quad sdk install <archive>", "Shipped"],
        ["Claude Code MCP integration (5 tools, 11 bundled skills)",
         "Shipped"],
        ["QUAD-Client provisioner (stdio-local / stdio-ssh / sse-http)",
         "Shipped"],
        ["bootstrap.ps1 (idempotent VS 2022 + Py 3.10 + redist + dlc_utils patch on ARM64 Windows)",
         "Shipped"],
        ["install.sh / install-client.sh with seamless detection",
         "Shipped"],
        ["Model registry: 7 entries (Plan 1 vision: 3 URL-fetchable; Plan 2 LLM: 2 env-var; Plan 4 ASR: 2 URL-fetchable)",
         "Shipped"],
    ])

    h2(doc, "Documentation")
    table(doc, ["Asset", "Status"], [
        ["Server Guide + Client Guide (auto-generated docx)",                  "Shipped"],
        ["Executive Pitch + User Journey PPTs",                                "Shipped"],
        ["Design Document QUAD Agent + Platform Design v2",                    "Shipped (with measurement-fidelity appendix)"],
        ["IoT Dependencies catalogue (111 components, 14 categories)",         "Shipped"],
        ["6 NPU Implementation Plans + 2 audience-perspective walkthroughs",   "Shipped"],
        ["USAGE_GUIDE, PRD, Sample App Prompts",                               "Shipped"],
        ["QUAD-Academy curriculum (12 modules + capstone)",                    "v0.1 (this snapshot)"],
        ["API reference (Sphinx auto-gen)",                                    "Not started"],
    ])

    h2(doc, "Testing")
    table(doc, ["Item", "Status"], [
        ["Unit + integration tests (full suite)",         "2000+ passing"],
        ["Adapter / parser / profiler / registry / sdk_patch tests", "All green"],
        ["Mock-mode e2e (tests/e2e/test_real_sdk_e2e.py)", "Green"],
        ["Real-hw smoke (projects/real_hw/smoke.py)",     "5/6 OK on X Elite"],
        ["Real-hw acceptance gates (@pytest.mark.real_hw)","Skipped — need real .dlc"],
        ["Perf regression CI",                            "Not yet wired"],
    ])


def section_pending(doc):
    h1(doc, "Pending — Tier 1 (pre-launch blockers)")
    p(doc, "These items gate a credible v1.0 commercial launch. Effort "
           "estimates assume one engineer focused on the item; some are "
           "blocked on resources outside the codebase.")

    table(doc, ["#", "Item", "Effort", "Blockers"], [
        ["1", "End-to-end validation on 10+ production models",
         "4–6 weeks",
         "QAIRT 2.46 internal bug we hit on MobileNetV2-12; need more test SoCs"],
        ["2", "Self-hosted X Elite perf regression CI",
         "1–2 weeks",
         "Dedicated CI runner"],
        ["3", "AIMET INT8 / INT4 accuracy regression loop",
         "3–4 weeks",
         "Calibration + reference accuracy datasets"],
        ["4", "PyPI publish under quad-agent",
         "1 week (blocked on Python 3.11+ wheel support)",
         "QAIRT must drop python310-only constraint"],
        ["5", "License finalisation + SPDX headers",
         "2 days",
         "Legal sign-off"],
        ["6", "API v1.0 freeze + deprecation policy",
         "1 week + ongoing governance",
         "—"],
        ["7", "Security review (template-injection, model signing, SBOM, dependency audit)",
         "2–3 weeks",
         "Reviewer resource"],
        ["8", "Cross-SoC validation: 8 Elite (Mobile), QCS6490 (IoT)",
         "4 weeks",
         "Physical hardware"],
    ])

    h2(doc, "Pending — Tier 2 (production-ops)")
    table(doc, ["Item", "Effort"], [
        ["QUAD Serve auth + rate limiting + remote model loading (S3 / Azure Blob)", "3–4 weeks"],
        ["OpenTelemetry tracing wired beyond just dep declarations",                  "1 week"],
        ["K8s manifests + Helm chart",                                                "2 weeks"],
        ["Per-SDK-version compatibility matrix CI",                                   "1–2 weeks"],
        ["QAIRT SDK bug-escalation workflow with Qualcomm",                           "Ongoing"],
        ["Sample-app gallery: 1 vision + 1 LLM + 1 ASR shipped with real models",     "≈ 6 weeks each"],
        ["Tutorial content: 5-min quickstart screencast + long-form tutorials",       "2–3 weeks"],
        ["Customer onboarding playbook",                                              "1 week"],
        ["Public model zoo (downloadable bundles)",                                   "2 weeks"],
    ])

    h2(doc, "Pending — Tier 3 (community / Phase H — currently 0%)")
    table(doc, ["Item", "Effort"], [
        ["Make repos public + CONTRIBUTING.md + Code of Conduct + SECURITY.md",     "1 week"],
        ["QUAD Academy v0.1 → v1.0: lecture content, recorded screencasts, MCQ bank → 200 questions, devcontainer / Codespaces image, exam platform integration", "≈ 3 months"],
        ["Discord / Stack Overflow tag presence",                                   "Ongoing"],
        ["Hackathon program",                                                       "1–2 quarters"],
        ["Public roadmap (GitHub Projects)",                                        "1 week"],
    ])

    h2(doc, "Pending — Tier 4 (compliance / operations)")
    table(doc, ["Item", "When"], [
        ["GDPR / CCPA stance for telemetry data",                            "Before public launch"],
        ["Export-control review (QAIRT redistribution constraints)",         "Before public launch"],
        ["Bug bounty / Vulnerability Disclosure Programme",                  "Post-public launch"],
        ["24/7 oncall infrastructure",                                       "Once paying customers"],
        ["Conformance to Qualcomm partner-program requirements",             "TBD"],
    ])


def section_can_be_done_now(doc):
    h1(doc, "What can ship next without external resources")
    p(doc, "These items are unblocked and could be picked up in the next "
           "engineering sprint without waiting for hardware, legal, or "
           "external Qualcomm action.")
    bullets(doc, [
        "License + SPDX header pass — apply Apache-2.0 / proprietary per file; drop a LICENSE",
        "SECURITY.md, CONTRIBUTING.md, Code of Conduct in all three repos",
        "Pre-commit tightening: bandit, pip-audit, ruff security rules",
        "SBOM generation script (Syft / CycloneDX) wired into install.sh",
        "API v1.0 freeze pass — mark public surfaces, add @deprecated decorator",
        "PyPI dry-run publish (python -m build + twine check)",
        "Public-repo prep: issue + PR templates, CODEOWNERS",
        "Customer-onboarding playbook (docs/ONBOARDING.md)",
        "QAIRT bug reproducer script for the MobileNetV2-12 broadcast issue (ready to file with Qualcomm)",
        "QUAD-Academy v0.5: write out remaining lab scaffolds (Modules 3-11), expand MCQ bank to 200 questions",
    ])

    h2(doc, "Not autonomously shippable (resource-gated)")
    bullets(doc, [
        "Buy / borrow more Snapdragon SoCs (X2 Elite, 8 Elite, QCS6490, RB3 Gen 2)",
        "Establish a Qualcomm SDK bug-escalation channel",
        "Run a persistent self-hosted CI runner",
        "Publish to PyPI under quad-agent (needs PyPI credentials + license decision)",
        "Make GitHub repos public (business decision)",
        "Build training videos and certification platform integration",
    ])


def section_recommendation(doc):
    h1(doc, "Recommendation")
    p(doc,
      "If the goal is a credible v1.0 launch in 8–12 weeks, focus on "
      "Tier-1 items 1, 2, 4, 5, 6, 7 in that order. Items 3 and 8 (perf "
      "and accuracy CI) can run in parallel and produce the evidence "
      "customers will demand.")
    p(doc,
      "If the goal is a community / OSS launch first (cheaper, faster), "
      "tackle Tier-3 community items plus Tier-1 items 5, 6, 7 — typically "
      "4–6 weeks of work — and let real-hardware validation grow from "
      "external contributors.")
    callout(doc,
            "QUAD Academy v0.1 ships alongside this status snapshot. "
            "The v0.5 → v1.0 progression of the Academy is the lowest-"
            "risk Tier-3 item and the highest-leverage one for "
            "community building.", "tip")

    h2(doc, "Reference commits this snapshot is based on")
    bullets(doc, [
        "QUAD: d83617c (Seamless QAIRT bring-up on Snapdragon X Elite: VS 2022 + Py 3.10 + dlc_utils patch)",
        "QUAD-Client: 060ba7c (Refresh Client Guide with measurement_notes + quad models FAQ)",
        "projects: 13782fd (real_hw/smoke.py passes input_name + input_dimensions for dynamic-shape ONNX)",
        "QUAD-Academy: v0.1 (this snapshot)",
    ])


def main() -> None:
    doc = Document()
    setup_styles(doc)

    title_page(doc)
    section_sources_of_truth(doc)
    doc.add_page_break()
    section_supported(doc)
    doc.add_page_break()
    section_pending(doc)
    doc.add_page_break()
    section_can_be_done_now(doc)
    doc.add_page_break()
    section_recommendation(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
