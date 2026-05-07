#!/usr/bin/env python3
"""Generate QUAD Sample App Prompts.docx"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

# ── Colour palette ────────────────────────────────────────────────────────────
QUALCOMM_RED  = RGBColor(0xCC, 0x00, 0x00)
DARK_GRAY     = RGBColor(0x2B, 0x2B, 0x2B)
MID_GRAY      = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY    = RGBColor(0xF4, 0xF4, 0xF4)
CODE_BG       = RGBColor(0xF0, 0xF0, 0xF0)
BLUE_ACCENT   = RGBColor(0x00, 0x5B, 0xB5)
GREEN_ACCENT  = RGBColor(0x1A, 0x7A, 0x3C)

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)

def add_heading(text, level=1, color=QUALCOMM_RED):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_para(text, bold=False, italic=False, color=DARK_GRAY, size=10.5,
             space_before=0, space_after=5, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    return p

def add_prompt_box(prompt_text, label="Prompt"):
    """Render a shaded prompt block."""
    # Label line
    p_label = doc.add_paragraph()
    p_label.paragraph_format.space_before = Pt(6)
    p_label.paragraph_format.space_after  = Pt(1)
    p_label.paragraph_format.left_indent  = Inches(0.2)
    r = p_label.add_run(f"💬  {label}")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = MID_GRAY

    # Prompt text in a shaded table (simulates a code/chat block)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, RGBColor(0xEF, 0xF4, 0xFF))
    cell.width = Inches(6.3)

    # Border colour tweak
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '3B6EC8')
        tcBorders.append(border)
    tcPr.append(tcBorders)

    cp = cell.paragraphs[0]
    cp.paragraph_format.left_indent  = Inches(0.15)
    cp.paragraph_format.right_indent = Inches(0.15)
    cp.paragraph_format.space_before = Pt(6)
    cp.paragraph_format.space_after  = Pt(6)
    cr = cp.add_run(prompt_text)
    cr.font.size = Pt(10)
    cr.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6E)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_result_box(items):
    """Add a shaded 'What QUAD returns' box."""
    p_label = doc.add_paragraph()
    p_label.paragraph_format.space_before = Pt(2)
    p_label.paragraph_format.space_after  = Pt(1)
    p_label.paragraph_format.left_indent  = Inches(0.2)
    r = p_label.add_run("📦  What QUAD returns")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREEN_ACCENT

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, RGBColor(0xF0, 0xFA, 0xF2))

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '1A7A3C')
        tcBorders.append(border)
    tcPr.append(tcBorders)

    cp = cell.paragraphs[0]
    cp.paragraph_format.left_indent  = Inches(0.15)
    cp.paragraph_format.space_before = Pt(4)
    cp.paragraph_format.space_after  = Pt(4)
    for item in items:
        run = cp.add_run(f"• {item}\n")
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x1A)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

def add_tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(f"💡  Tip: {text}")
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x88, 0x66, 0x00)

def add_divider():
    p = doc.add_paragraph("─" * 72)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(8)

# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════

p_title = doc.add_paragraph()
p_title.paragraph_format.space_before = Pt(40)
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run("QUAD")
r.font.size  = Pt(52)
r.font.bold  = True
r.font.color.rgb = QUALCOMM_RED

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_sub.add_run("Qualcomm Unified Agent for Developers")
r.font.size  = Pt(18)
r.font.color.rgb = DARK_GRAY

doc.add_paragraph()

p_desc = doc.add_paragraph()
p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_desc.add_run("Sample App Prompts")
r.font.size  = Pt(22)
r.font.bold  = True
r.font.color.rgb = BLUE_ACCENT

doc.add_paragraph()

p_tagline = doc.add_paragraph()
p_tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_tagline.add_run(
    "Copy-paste prompts for Claude Code + QUAD MCP Server\n"
    "Works in mock mode — no hardware or SDK required"
)
r.font.size  = Pt(11)
r.font.italic = True
r.font.color.rgb = MID_GRAY

doc.add_paragraph()
doc.add_paragraph()

# Quick start box
tbl = doc.add_table(rows=1, cols=1)
tbl.style = 'Table Grid'
cell = tbl.cell(0, 0)
set_cell_bg(cell, RGBColor(0xFF, 0xF8, 0xEE))
cp = cell.paragraphs[0]
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp.paragraph_format.space_before = Pt(8)
cp.paragraph_format.space_after  = Pt(8)
r = cp.add_run("Quick Start:  git clone <repo> && cd QUAD && ./setup.sh && source .venv/bin/activate")
r.font.size = Pt(10)
r.font.bold = True
r.font.color.rgb = RGBColor(0x66, 0x33, 0x00)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ════════════════════════════════════════════════════════════════════════════

add_heading("Contents", level=1, color=DARK_GRAY)
toc = [
    ("1", "Getting Started",                     "How QUAD works with Claude Code"),
    ("2", "Hardware Detection",                  "Discover chipset, CPU, GPU, NPU specs"),
    ("3", "Model Conversion",                    "Convert ONNX/PyTorch/TF to DLC"),
    ("4", "Profiling",                           "Latency, power, HTP linting, QHAS"),
    ("5", "Workload Orchestration",              "Allocate layers across CPU/GPU/NPU"),
    ("6", "Code Generation",                     "Generate C++, Python, Kotlin"),
    ("7", "Full End-to-End Workflows",           "Multi-step pipeline prompts"),
    ("8", "Multi-Model & Platform Comparison",   "Compare models side-by-side"),
    ("9", "Optimization Workflows",              "Find and fix HTP bottlenecks"),
    ("10", "Advanced Scenarios",                 "Power modes, quantization, QHAS"),
]

for num, title, desc in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.3)
    r1 = p.add_run(f"{num}.  {title}")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = BLUE_ACCENT
    r2 = p.add_run(f"  —  {desc}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = MID_GRAY

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. GETTING STARTED
# ════════════════════════════════════════════════════════════════════════════

add_heading("1. Getting Started")
add_para(
    "QUAD exposes 5 MCP tools that Claude Code can call: hardware_detect, convert_model, "
    "profile_workload, orchestrate_workload, and generate_code. Once the server is running, "
    "just describe what you want — Claude calls the right tools automatically.",
    size=10.5, color=MID_GRAY, space_after=8
)
add_para(
    "All prompts below work in mock mode — no Qualcomm hardware or SDK installation needed.",
    bold=True, size=10.5, color=GREEN_ACCENT, space_after=10
)

add_para("Start the MCP server:", bold=True, size=10, space_after=2)
add_prompt_box("quad-server --config quad.toml", label="Terminal")

add_para("Open this project in Claude Code — QUAD tools are auto-registered.", size=10, color=MID_GRAY, space_after=10)

add_tip("Use 'mock mode' for all development and demos. Switch to 'real' mode only when the QNN/SNPE SDK is installed.")
add_divider()

# ════════════════════════════════════════════════════════════════════════════
# 2. HARDWARE DETECTION
# ════════════════════════════════════════════════════════════════════════════

add_heading("2. Hardware Detection")
add_para("Detect the chipset, CPU, GPU, and NPU topology for any Qualcomm platform.", color=MID_GRAY, space_after=8)

add_para("Detect Windows AI PC (Snapdragon X Elite):", bold=True, size=10)
add_prompt_box(
    "Use QUAD to detect the hardware on my Windows AI PC. "
    "Show the chipset name, CPU cores, GPU TFLOPS, and NPU TOPS."
)
add_result_box([
    "Chipset: Snapdragon X Elite X1E-80-100",
    "CPU: 12 × Oryon ARM64 @ 3.8 GHz",
    "GPU: Adreno X1-85 (4.6 TFLOPS)",
    "NPU: Hexagon NPU (45.0 TOPS)",
    "Available runtimes: cpu, gpu, npu",
])

add_para("Detect Android device (Snapdragon 8 Elite):", bold=True, size=10)
add_prompt_box(
    "Detect hardware for the Android platform using QUAD. "
    "What NPU is available and how many TOPS does it have?"
)
add_result_box([
    "Chipset: Snapdragon 8 Elite SM8750",
    "NPU: Hexagon NPU (HTP) — 48.0 TOPS",
    "GPU: Adreno 830 (5.0 TFLOPS)",
])

add_para("Detect Linux embedded (QCS2210):", bold=True, size=10)
add_prompt_box(
    "Use QUAD hardware_detect for the linux platform. "
    "Show all compute resources available."
)
add_result_box([
    "Chipset: QCS2210 (Qualcomm Robotics RB1)",
    "CPU: 4 × Kryo ARM64 @ 2.0 GHz",
    "NPU: Hexagon DSP V66 (1.0 TOPS)",
    "RAM: 2 GB",
])
add_divider()

# ════════════════════════════════════════════════════════════════════════════
# 3. MODEL CONVERSION
# ════════════════════════════════════════════════════════════════════════════

add_heading("3. Model Conversion")
add_para("Convert ONNX, PyTorch, TensorFlow, or TFLite models to SNPE DLC or QNN format.", color=MID_GRAY, space_after=8)

add_para("Basic conversion — INT8 quantization:", bold=True, size=10)
add_prompt_box(
    "Use QUAD to convert resnet50.onnx to SNPE format with INT8 quantization for Windows. "
    "Show the output file, compression ratio, and any image format requirements."
)
add_result_box([
    "Output: resnet50.dlc",
    "Original: 25.0 MB  →  Converted: 6.2 MB  (4.0× compression)",
    "Quantization: int8 applied",
    "Image format note: NCHW → NHWC transpose needed before inference",
])

add_para("PyTorch model with BGR channel order (e.g. AlexNet):", bold=True, size=10)
add_prompt_box(
    "Convert alexnet.onnx to SNPE DLC. The model was trained with BGR channel order "
    "and PyTorch NCHW layout. Show me exactly how to prepare input images."
)
add_result_box([
    "channel_order: bgr — ensure input images are in BGR order",
    "input_layout: nchw — transpose: np.transpose(img, (0,2,3,1)) → NHWC",
    "Output: alexnet.dlc",
])

add_para("MobilenetSSD object detection model:", bold=True, size=10)
add_prompt_box(
    "Convert ssd_mobilenet_v2_quantized_300x300.pb (TensorFlow frozen graph) "
    "to SNPE DLC for GPU inference. Show any special requirements for this model."
)
add_result_box([
    "Conversion notes: MobilenetSSD requires allow_unconsumed_nodes=True",
    "DetectionOutput layer: CPU only — enable CPU fallback for GPU/DSP",
    "Output nodes: detection_classes, detection_boxes, detection_scores",
])

add_tip("Always check conversion_notes in the result — QUAD surfaces model-specific tips automatically.")
add_divider()

# ════════════════════════════════════════════════════════════════════════════
# 4. PROFILING
# ════════════════════════════════════════════════════════════════════════════

add_heading("4. Profiling")
add_para("Profile inference performance at multiple levels: timing, HTP cycle counts, or full QHAS analysis.", color=MID_GRAY, space_after=8)

add_para("Standard detailed profiling:", bold=True, size=10)
add_prompt_box(
    "Profile resnet50.dlc on the Windows NPU using QUAD. "
    "Show mean latency, p95 latency, throughput, power consumption, and memory usage."
)
add_result_box([
    "Latency: mean=5.0ms  p95=7.0ms  p99=9.0ms",
    "Throughput: 200 FPS",
    "Power: 2000 mW",
    "Memory: peak=45 MB  avg=38 MB",
    "NPU utilization: 89%",
])

add_para("HTP Linting profile — cycle-based bottleneck analysis:", bold=True, size=10)
add_prompt_box(
    "Use QUAD to profile mobilenetv2.dlc with HTP linting mode. "
    "Identify the top bottleneck ops and explain how to fix them."
)
add_result_box([
    "Total cycles: 3,753,939",
    "Bottleneck: sub_op (2,165,162 cycles, 21% overlap) ⚠️",
    "Hint: Replace Sub with Conv for better HTP parallelism (~68% improvement)",
    "Well-parallelized: add_op (92% overlap) ✅",
])

add_para("QHAS profiling — full QNN HTP Analysis Summary:", bold=True, size=10)
add_prompt_box(
    "Run QHAS profiling on yolov8n.dlc using QUAD. "
    "What chrometrace file does it generate and how do I view it?"
)
add_result_box([
    "Runs 3-step pipeline: graph-prepare → net-run → profile-viewer",
    "Generates: chrometrace.json (open in chrome://tracing)",
    "Also produces: _htp.json with op-by-op topology",
])

add_tip("Use 'detailed' for latency/power numbers. Use 'linting' to find HTP bottlenecks. Use 'qhas' for deep analysis with chrome tracing.")
add_divider()

# ════════════════════════════════════════════════════════════════════════════
# 5. WORKLOAD ORCHESTRATION
# ════════════════════════════════════════════════════════════════════════════

add_heading("5. Workload Orchestration")
add_para("Automatically allocate model layers across CPU, GPU, and NPU for a given power budget.", color=MID_GRAY, space_after=8)

add_para("Performance mode — maximize NPU usage:", bold=True, size=10)
add_prompt_box(
    "Use QUAD to orchestrate resnet50.dlc on Windows in performance mode. "
    "Show which layers go to NPU vs CPU and the projected latency."
)
add_result_box([
    "NPU utilization: 70%  |  CPU: 30%  |  GPU: 0%",
    "Projected latency: 2.55 ms",
    "Projected power: 1420 mW",
    "CPU fallback ops: BatchNorm layers (unsupported on HTP)",
])

add_para("Compare all three power modes:", bold=True, size=10)
add_prompt_box(
    "Using QUAD, compare performance, balanced, and efficiency power modes "
    "for mobilenetv2.dlc on Android. Show latency and power for each."
)
add_result_box([
    "performance → 2.55ms  1420mW  NPU=70%",
    "balanced    → 2.55ms  1420mW  NPU=70%",
    "efficiency  → 5.00ms  1000mW  NPU=0%  (CPU only)",
])

add_para("Low-power embedded (QCS2210, 3W budget):", bold=True, size=10)
add_prompt_box(
    "Orchestrate yolov8n.dlc for the linux platform in efficiency mode. "
    "The device has a 3W power budget. What's the recommended allocation?"
)
add_divider()

# ════════════════════════════════════════════════════════════════════════════
# 6. CODE GENERATION
# ════════════════════════════════════════════════════════════════════════════

add_heading("6. Code Generation")
add_para("Generate production-ready inference code for C++, Python, Kotlin, or Arduino.", color=MID_GRAY, space_after=8)

add_para("C++ inference app for Windows AI PC:", bold=True, size=10)
add_prompt_box(
    "Use QUAD to generate a C++ inference application for resnet50.dlc "
    "on Windows using the QNN SDK. Include the CMakeLists.txt build file."
)
add_result_box([
    "inference.cpp — full QNN initialization, graph execute, output read",
    "CMakeLists.txt — build config targeting Windows ARM64",
    "Build: mkdir build && cd build && cmake .. && cmake --build .",
])

add_para("Python inference script for Android:", bold=True, size=10)
add_prompt_box(
    "Generate Python inference code for mobilenetv2.dlc on Android "
    "using SNPE. I need to run it via ADB."
)

add_para("Kotlin Android AAR for mobile app:", bold=True, size=10)
add_prompt_box(
    "Use QUAD generate_code for the Android platform with Kotlin language "
    "for yolov8n.dlc. Generate an Android inference module I can add to my app."
)

add_para("Arduino sketch for embedded Linux:", bold=True, size=10)
add_prompt_box(
    "Generate an Arduino-compatible inference sketch for the linux platform "
    "using SNPE SDK for a simple classification model."
)
add_divider()

# ════════════════════════════════════════════════════════════════════════════
# 7. FULL END-TO-END WORKFLOWS
# ════════════════════════════════════════════════════════════════════════════

add_heading("7. Full End-to-End Workflows")
add_para("Use multiple QUAD tools together in a single conversation flow.", color=MID_GRAY, space_after=8)

add_para("ResNet-50 on Windows AI PC — full pipeline:", bold=True, size=10)
add_prompt_box(
    "Build a complete inference pipeline for resnet50.onnx on Windows AI PC using QUAD:\n"
    "1. Detect the hardware\n"
    "2. Convert to SNPE DLC with INT8 quantization\n"
    "3. Profile with detailed timing\n"
    "4. Profile with HTP linting to find bottlenecks\n"
    "5. Orchestrate in balanced mode\n"
    "6. Generate C++ inference code\n"
    "Save everything to examples/resnet50_windows.py"
)

add_para("MobileNetV2 on Android — TTFI under 5 seconds:", bold=True, size=10)
add_prompt_box(
    "Using QUAD, run the full workflow for mobilenetv2.onnx on Android. "
    "I need Time-to-First-Inference under 5 seconds. "
    "Use INT8 quantization and balanced power mode. "
    "Generate Kotlin code at the end."
)

add_para("YOLOv8 object detection on all 3 platforms:", bold=True, size=10)
add_prompt_box(
    "Use QUAD to run the full pipeline for yolov8n.onnx across all three platforms: "
    "Windows AI PC, Android, and Linux/QCS2210. "
    "Compare the latency and NPU utilization for each. "
    "Generate Python code for each platform."
)
add_divider()

# ════════════════════════════════════════════════════════════════════════════
# 8. MULTI-MODEL & PLATFORM COMPARISON
# ════════════════════════════════════════════════════════════════════════════

add_heading("8. Multi-Model & Platform Comparison")
add_para("Use QUAD to benchmark and compare multiple models side-by-side.", color=MID_GRAY, space_after=8)

add_para("Latency comparison across models:", bold=True, size=10)
add_prompt_box(
    "Use QUAD to compare inference latency for three models on Windows AI PC: "
    "mobilenetv2.onnx, resnet50.onnx, and yolov8n.onnx. "
    "Use INT8 quantization and NPU runtime for all. "
    "Present results in a table showing latency, throughput, and power."
)

add_para("Platform suitability analysis:", bold=True, size=10)
add_prompt_box(
    "I have mobilenetv2.onnx. Use QUAD to help me decide which Qualcomm platform "
    "is best for my use case — Windows AI PC, Android, or embedded Linux. "
    "Compare NPU TOPS, expected latency, and power consumption."
)

add_para("SDK comparison — QNN vs SNPE:", bold=True, size=10)
add_prompt_box(
    "Use QUAD to convert resnet50.onnx twice — once to QNN format and once to SNPE DLC. "
    "Compare the output file sizes and tell me which SDK is better for my Windows AI PC use case."
)
add_divider()

# ════════════════════════════════════════════════════════════════════════════
# 9. OPTIMIZATION WORKFLOWS
# ════════════════════════════════════════════════════════════════════════════

add_heading("9. Optimization Workflows")
add_para("Use QUAD linting and architecture analysis to identify and fix performance bottlenecks.", color=MID_GRAY, space_after=8)

add_para("Find and fix HTP bottlenecks:", bold=True, size=10)
add_prompt_box(
    "Profile my_model.dlc with QUAD linting mode. "
    "For each bottleneck op found (low overlap %), explain why it's slow "
    "and what op substitution I should make to improve HTP parallelism."
)
add_result_box([
    "Sub op → replace with Conv2D (designed weights): −68% cycles",
    "Div op → replace with Mul (reciprocal): −65% cycles",
    "PReLU → replace with ReLU: better HTP parallelism",
])

add_para("Architecture checker analysis:", bold=True, size=10)
add_prompt_box(
    "Run QUAD's architecture checker workflow on my_model.dlc. "
    "Show all HTP performance issues and apply the recommended modifications."
)
add_result_box([
    "Issue: 16-bit activation → recommend 8-bit for better memory efficiency",
    "Issue: Conv channel < 32 → increase to 32+ for better HTP utilization",
    "Issue: ElementWiseDivide → replace with ElementWiseMultiply",
    "Modifications applied via --modify apply=elwisediv",
])

add_para("Quantization impact analysis:", bold=True, size=10)
add_prompt_box(
    "Use QUAD to compare the performance impact of FP32 vs INT8 vs INT4 quantization "
    "for resnet50.onnx on the Snapdragon X Elite NPU. "
    "Show latency, memory, and compression ratio for each."
)
add_divider()

# ════════════════════════════════════════════════════════════════════════════
# 10. ADVANCED SCENARIOS
# ════════════════════════════════════════════════════════════════════════════

add_heading("10. Advanced Scenarios")

add_para("QHAS chrometrace generation:", bold=True, size=10)
add_prompt_box(
    "Use QUAD to run QHAS profiling on mobilenetv2.dlc for SM8750. "
    "Generate a chrometrace JSON I can open in chrome://tracing. "
    "Enable all features: input/output flow events, runtrace, memory_info."
)

add_para("MobilenetSSD benchmarking setup:", bold=True, size=10)
add_prompt_box(
    "Use QUAD to set up a MobilenetSSD benchmark run. "
    "Generate the benchmark config JSON, input list file with the correct "
    "output layer header, and the snpe_bench.py command to run it."
)
add_result_box([
    "mobilenetssd_bench.json with CpuFallback=true and BufferTypes=[ub_float, ub_tf8]",
    "imagelist.txt with header: #Postprocessor/BatchMultiClassNonMaxSuppression add_6",
    "Command: python3 snpe_bench.py -c mobilenetssd_bench.json -a --generate_json",
])

add_para("Custom UDO (User-Defined Operation) scaffolding:", bold=True, size=10)
add_prompt_box(
    "I need a custom Softmax operation for the HTP backend. "
    "Use QUAD to generate the UDO package scaffold — config JSON, "
    "DSP implementation stub, and registration library."
)

add_para("Multi-input model with named tensors:", bold=True, size=10)
add_prompt_box(
    "Convert my two-input ONNX model to SNPE DLC. "
    "The inputs are: 'image' with shape 1,3,224,224 (NCHW, RGB) "
    "and 'mask' with shape 1,1,224,224. "
    "The output tensor is 'predictions'. "
    "Show the exact conversion command."
)

add_para("Accuracy debugging workflow:", bold=True, size=10)
add_prompt_box(
    "My SNPE model gives different results from the original ONNX model. "
    "Use QUAD to set up an accuracy debugger run: "
    "run framework_runner on the ONNX, run inference_engine on the DLC, "
    "and compare using CosineSimilarity. Show the verification command."
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# QUICK REFERENCE TABLE
# ════════════════════════════════════════════════════════════════════════════

add_heading("Quick Reference — Tool Parameters")

headers = ["Parameter", "Options", "Notes"]
rows = [
    ["platform",        "windows  android  linux",              "Selects mock device profile"],
    ["source_format",   "onnx  pytorch  tensorflow  tflite",    "Model framework"],
    ["target_sdk",      "snpe  qnn",                            "Output format"],
    ["quantization",    "fp32  int8  int4",                     "int8 recommended for NPU"],
    ["input_layout",    "nchw  nhwc  auto",                     "PyTorch = nchw"],
    ["channel_order",   "rgb  bgr  auto",                       "Caffe/legacy = bgr"],
    ["profiling_level", "basic  detailed  linting  qhas",       "linting = HTP cycles"],
    ["power_mode",      "performance  balanced  efficiency",    "Controls NPU vs CPU split"],
    ["language",        "cpp  python  kotlin  arduino_sketch",  "Code gen target"],
    ["runtime",         "cpu  gpu  npu  auto",                  "Inference runtime"],
]

tbl = doc.add_table(rows=1 + len(rows), cols=3)
tbl.style = 'Table Grid'

# Header row
hdr_cells = tbl.rows[0].cells
for i, h in enumerate(headers):
    set_cell_bg(hdr_cells[i], QUALCOMM_RED)
    p = hdr_cells[i].paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(10)

# Data rows
for ri, row_data in enumerate(rows):
    row_cells = tbl.rows[ri + 1].cells
    bg = RGBColor(0xF8, 0xF8, 0xF8) if ri % 2 == 0 else RGBColor(0xFF, 0xFF, 0xFF)
    for ci, val in enumerate(row_data):
        set_cell_bg(row_cells[ci], bg)
        p = row_cells[ci].paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(9.5)
        if ci == 0:
            r.bold = True
            r.font.color.rgb = BLUE_ACCENT
        elif ci == 1:
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x66)

doc.add_paragraph()

# ── Footer note ───────────────────────────────────────────────────────────────
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.paragraph_format.space_before = Pt(20)
r = p_footer.add_run(
    "QUAD — Qualcomm Unified Agent for Developers  |  "
    "All prompts work in mock mode without hardware or SDK  |  "
    "github.qualcomm.com/pavanr/QUAD"
)
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = "docs/QUAD_Sample_App_Prompts.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
